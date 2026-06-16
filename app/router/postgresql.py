# -*- coding: utf-8 -*-
"""
项目与文档 CRUD 路由。

提供项目、文档、关系、分析结果的增删改查接口，
包含文档预览（含高亮）、重复检查、形式审查、人员复用检查、错别字检查等功能。
"""

import base64
import json
import io
import os
import hashlib
import logging
import re
import tempfile
from copy import deepcopy
from datetime import datetime
from typing import Any, Literal, Optional
from urllib.parse import quote

from fastapi import APIRouter, Body, Depends, File, Form, HTTPException, Query, Response, UploadFile
from fastapi.responses import JSONResponse, RedirectResponse
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from psycopg2 import Error as PsycopgError
from starlette.concurrency import run_in_threadpool

from app.config.settings import settings
from app.core.document_types import (
    BID_DOCUMENT_TYPES,
    DOCUMENT_TYPE_BUSINESS_BID,
    DOCUMENT_TYPE_TENDER,
    DOCUMENT_TYPE_TECHNICAL_BID,
    DocumentType,
)
from app.router.dependencies import (
    RecognitionOptions,
    get_bid_document_review_service,
    get_cache_service,
    get_db_service,
    get_duplicate_check_service,
    get_form_recognition_options,
    get_oss_service,
    get_text_analysis_service,
)
from app.router.uploaded_json_support import (
    build_uploaded_project_document_records,
    load_uploaded_bid_json_documents,
    persist_uploaded_json_project_documents,
    read_uploaded_json_file,
)
from app.schemas.postgresql import (
    BusinessBidManualReviewInputsRequest,
    DuplicateCheckScope,
    DocumentReviewContentUpdateRequest,
    DocumentPreviewRequest,
    DocumentUpdateRequest,
    IdentifierBatchDeleteRequest,
    PersonnelReuseCheckRequest,
    PersonnelReuseDraftRequest,
    ProjectBindDocumentsRequest,
    ProjectCreateRequest,
    ProjectDuplicateCheckRequest,
    ProjectManualReviewRerunRequest,
    ProjectManualReviewResultInputsRequest,
    ProjectReportExportRequest,
    ProjectResultUpdateRequest,
    ProjectResultUpsertRequest,
    ProjectRelationUpdateRequest,
    ProjectWorkflowScopeRequest,
    RelationBatchDeleteRequest,
    ProjectUpdateRequest,
)
from app.service.analysis import BidDocumentReviewService, DuplicateCheckService
from app.service.cache_service import CacheUnavailableError, RedisCacheService
from app.service.analysis.duplicate_merge import (
    DOC_TYPE_BY_MERGED_RESULT_KEY,
    MERGE_STRATEGY,
    MERGED_RESULT_KEY_BY_DOC_TYPE,
    RAW_RESULT_KEY_BY_DOC_TYPE,
    build_duplicate_merge_results,
)
from app.service.analysis.manual_review.business_bid_format import (
    BUSINESS_FORMAT_RESULT_KEY,
    _apply_manual_business_review_inputs,
    _build_business_format_editable_items,
    _business_manual_payload_for_project,
    _business_review_from_record,
    _save_business_manual_inputs,
)
from app.service.analysis.project_input_loader import ProjectAnalysisInputLoader
from app.service.manual_review_state import (
    MANUAL_REVIEW_RESULTS_KEY,
    WORKFLOW_SCOPE_KEY,
    display_result_view,
    manual_review_results_from_record,
    raw_result_view,
)
from app.service.manual_review.working_copy import (
    DocumentWorkingCopyService,
)
from app.service.workflow_scope import (
    EXCLUDED_BIDDERS_KEY,
    filter_project_payload,
    normalize_workflow_scope,
    workflow_scope_from_result_record,
)
from app.service.analysis.unified import UnifiedBusinessReviewService
from app.service.document_ingest_service import normalize_file_url, upload_extract_and_create_document
from app.service.minio_service import MinioService
from app.service.postgresql_service import PostgreSQLService
from app.service.project_runtime import (
    active_project_runtime_identifiers,
    cancel_project_runtime,
)

router = APIRouter()
logger = logging.getLogger(__name__)

_FINE_OCR_PREVIEW_STRATEGY_VERSION = "fine-ocr-v1"
# 用于从高亮短语中提取有效 token 的正则
_HIGHLIGHT_TOKEN_PATTERN = re.compile(r"[\u4e00-\u9fffA-Za-z0-9]+")


def _set_cache_header(response: Optional[Response], status: str) -> None:
    if response is not None:
        response.headers["X-XTJS-Cache"] = status


def _cache_get_or_set_payload(
    *,
    cache_service: RedisCacheService,
    cache_key: str,
    ttl_seconds: int,
    response: Optional[Response],
    factory,
):
    if not settings.XTJS_CACHE_ENABLED:
        _set_cache_header(response, "disabled")
        return factory()
    try:
        payload, cache_status = cache_service.get_or_set_json(cache_key, ttl_seconds, factory)
        _set_cache_header(response, cache_status)
        return payload
    except CacheUnavailableError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc


def _invalidate_project_cache_or_error(cache_service: RedisCacheService, identifier_id: Optional[str] = None) -> None:
    if not settings.XTJS_CACHE_ENABLED:
        return
    try:
        cache_service.invalidate_project(identifier_id)
    except CacheUnavailableError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc


def _invalidate_document_preview_cache_or_error(cache_service: RedisCacheService, document_id: Optional[str] = None) -> None:
    if not settings.XTJS_CACHE_ENABLED:
        return
    try:
        cache_service.invalidate_document_preview(document_id)
    except CacheUnavailableError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc


def _invalidate_project_cache_by_identifier(identifier_id: Optional[str] = None) -> None:
    _invalidate_project_cache_or_error(get_cache_service(), identifier_id)


def _is_result_key_visible(result_key: str) -> bool:
    if result_key == MANUAL_REVIEW_RESULTS_KEY:
        return False
    if result_key in {"business_itemized_duplicate_check", "bid_response_duplicate_check"}:
        return False
    if result_key == "typo_check":
        # 独立错别字检查已下线，历史结果不再展示。
        return False
    return True


def _filter_visible_result_keys(results: dict[str, Any]) -> dict[str, Any]:
    return {
        key: value
        for key, value in dict(results or {}).items()
        if _is_result_key_visible(str(key))
    }


_PROJECT_SERVICE_RESULT_KEYS = {
    "business_bid_format_review": UnifiedBusinessReviewService.BUSINESS_RESULT_KEY,
    "deviation_check": "deviation_check",
    "business_bid_duplicate_check": "business_bid_duplicate_check",
    "technical_bid_duplicate_check": "technical_bid_duplicate_check",
    "personnel_reuse_check": "personnel_reuse_check",
}


# 将查重范围枚举转换为文档类型列表
def _normalize_selected_services(services: list[str] | None) -> list[str]:
    normalized: list[str] = []
    seen: set[str] = set()
    for service_name in services or []:
        token = str(service_name or "").strip()
        if not token or token in seen:
            continue
        seen.add(token)
        normalized.append(token)
    return normalized


def _http_exception_detail_to_text(detail: Any) -> str:
    if isinstance(detail, str):
        return detail
    return json.dumps(detail, ensure_ascii=False)


def _document_types_from_scope(scope: DuplicateCheckScope) -> Optional[list[str]]:
    if scope == DuplicateCheckScope.ALL:
        return None
    return [scope.value]


def _refresh_project_or_404(db_service: PostgreSQLService, identifier_id: str) -> dict[str, Any]:
    """先重算项目 OCR 状态，再返回最新项目记录。"""
    project = db_service.refresh_project_parsing_status(identifier_id)
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")
    return project


def _ensure_project_analysis_status(
    project: dict[str, Any],
    *,
    required_status: int,
    analysis_name: str,
) -> dict[str, Any]:
    """分析前先校验项目 OCR 阶段是否达标。"""
    current_status = int(project.get("parsing_status") or 0)
    if PostgreSQLService.parsing_status_reached(current_status, required_status):
        return project
    current_text = PostgreSQLService.get_parsing_status_text(current_status)
    required_text = PostgreSQLService.get_parsing_status_text(required_status)
    raise HTTPException(
        status_code=409,
        detail=(
            f"{analysis_name} 需要项目 OCR 状态至少达到 {required_status}（{required_text}），"
            f"当前为 {current_status}（{current_text}）。"
        ),
    )


def _ensure_project_ocr_idle(project: dict[str, Any], *, analysis_name: str) -> None:
    """OCR 后台执行时，先阻止会落库的项目级分析抢占运行资源。"""
    project_identifier = str(project.get("identifier_id") or "").strip()
    active_identifiers = active_project_runtime_identifiers()
    if not active_identifiers:
        return
    if project_identifier in active_identifiers:
        raise HTTPException(
            status_code=409,
            detail=(
                f"{analysis_name} 未启动：当前项目 OCR 仍在后台运行或排队。"
                "为避免业务分析读取半成品 OCR 并抢占运行资源，请等待 OCR 完成后重试；"
                f"可先查询 /api/postgresql/projects/{project_identifier}/ocr-status 查看进度。"
            ),
        )
    active_preview = "、".join(active_identifiers[:3])
    if len(active_identifiers) > 3:
        active_preview = f"{active_preview} 等 {len(active_identifiers)} 个项目"
    raise HTTPException(
        status_code=409,
        detail=(
            f"{analysis_name} 未启动：当前有其他项目 OCR 仍在后台运行或排队（{active_preview}）。"
            "为避免业务分析抢占 OCR 运行资源并影响其他项目 OCR，请等待 OCR 完成后重试；"
            "可先查询对应项目的 /api/postgresql/projects/{project_id}/ocr-status 查看进度。"
        ),
    )


def _required_parsing_status_for_duplicate_scope(scope: DuplicateCheckScope) -> int:
    # 综合查重只要带技术标范围，就必须等到技术标 OCR 完成。
    if scope == DuplicateCheckScope.BUSINESS_BID:
        return PostgreSQLService.PARSING_STATUS_BUSINESS_OCR_COMPLETED
    return PostgreSQLService.PARSING_STATUS_TECHNICAL_OCR_COMPLETED


def _required_parsing_status_for_document_types(document_types: Optional[list[str]]) -> int:
    normalized = {
        str(document_type or "").strip().lower()
        for document_type in (document_types or [])
        if str(document_type or "").strip()
    }
    if normalized and normalized <= {DOCUMENT_TYPE_BUSINESS_BID}:
        return PostgreSQLService.PARSING_STATUS_BUSINESS_OCR_COMPLETED
    return PostgreSQLService.PARSING_STATUS_TECHNICAL_OCR_COMPLETED


# 统一分页参数解析（兼容 page/limit 两种方式）
def _resolve_pagination(
    *,
    page: int,
    page_size: int,
    limit: Optional[int] = None,
    offset: Optional[int] = None,
) -> tuple[int, int]:
    if limit is not None or offset is not None:
        normalized_limit = max(1, min(int(limit or page_size), 200))
        normalized_offset = max(0, int(offset or 0))
        return normalized_limit, normalized_offset
    normalized_page_size = max(1, min(int(page_size), 200))
    normalized_page = max(1, int(page))
    return normalized_page_size, (normalized_page - 1) * normalized_page_size


# 根据文件名/URL 后缀判断文档源类型（pdf 或 image）
def _document_source_kind(document: dict) -> str:
    file_name = str(document.get("file_name") or "").strip()
    file_url = str(document.get("file_url") or "").strip()
    target = file_name or file_url
    suffix = os.path.splitext(target)[1].lower()
    if suffix == ".pdf":
        return "pdf"
    if suffix in {".png", ".jpg", ".jpeg"}:
        return "image"
    return "unknown"


# 从文档记录中解析 MinIO 的 bucket 和 object 名称
def _resolve_document_source_object(document: dict) -> tuple[str, str, str]:
    file_url = str(document.get("file_url") or "").strip()
    if not file_url:
        raise ValueError("document file_url is empty")

    if file_url.startswith("minio://"):
        bucket_name, object_name = MinioService.bucket_and_object_from_file_url(file_url)
    elif MinioService.is_presigned_url(file_url):
        bucket_name, object_name = MinioService.bucket_and_object_from_presigned_url(file_url)
    else:
        raise ValueError("document source file is not stored in MinIO")

    file_name = str(document.get("file_name") or object_name or "document").strip() or "document"
    return bucket_name, object_name, file_name


# 从 MinIO 下载文档原始字节
def _load_document_source_bytes(
    *,
    document: dict,
    oss_service: MinioService,
) -> tuple[bytes, str, str]:
    bucket_name, object_name, _file_name = _resolve_document_source_object(document)
    data, content_type = oss_service.get_object_bytes(object_name, bucket_name)
    return data, content_type, object_name


def _preview_document_version(document: dict) -> str:
    raw_version = "|".join(
        [
            str(document.get("update_time") or ""),
            str(document.get("file_url") or ""),
            str(document.get("source_file_hash") or ""),
        ]
    )
    return hashlib.sha1(raw_version.encode("utf-8")).hexdigest()


def _preview_cache_object_name(document: dict, version: str, page: int) -> str:
    document_id = str(document.get("identifier_id") or "").strip()
    prefix = str(settings.XTJS_CACHE_PREVIEW_OBJECT_PREFIX or "cache/previews").strip("/")
    return f"{prefix}/{document_id}/{version}/p{int(page)}.png"


def _raw_preview_payload_from_source(
    *,
    file_bytes: bytes,
    source_kind: str,
    page: int,
) -> dict[str, Any]:
    if page <= 0:
        raise ValueError("page must be greater than 0")

    if source_kind == "pdf":
        import fitz

        pdf = fitz.open(stream=file_bytes, filetype="pdf")
        try:
            page_count = int(pdf.page_count)
            if page > page_count:
                raise ValueError(f"page {page} is out of range, max page is {page_count}")
            pdf_page = pdf.load_page(page - 1)
            rect = pdf_page.rect
            pix = pdf_page.get_pixmap(matrix=fitz.Matrix(1.8, 1.8), alpha=False)
            return {
                "page": page,
                "page_count": page_count,
                "width": float(rect.width),
                "height": float(rect.height),
                "image_bytes": pix.tobytes("png"),
                "source_kind": "pdf",
            }
        finally:
            pdf.close()

    if source_kind == "image":
        if page != 1:
            raise ValueError("image document only supports page 1 preview")
        from PIL import Image

        with Image.open(io.BytesIO(file_bytes)) as image:
            rgb = image.convert("RGB")
            buffer = io.BytesIO()
            rgb.save(buffer, format="PNG")
            return {
                "page": 1,
                "page_count": 1,
                "width": int(rgb.width),
                "height": int(rgb.height),
                "image_bytes": buffer.getvalue(),
                "source_kind": "image",
            }

    raise ValueError("document source kind does not support preview")


def _raw_preview_response_payload(
    *,
    document: dict,
    page: int,
    image_bytes: bytes,
    meta: dict[str, Any],
) -> dict[str, Any]:
    return {
        "page": int(page),
        "page_count": int(meta.get("page_count") or 1),
        "width": float(meta.get("width") or 0),
        "height": float(meta.get("height") or 0),
        "image_data_url": _render_png_data_url(image_bytes),
        "source_kind": str(meta.get("source_kind") or "unknown"),
        "highlight_rect_count": 0,
        "highlight_applied": False,
        "highlight_coordinate_space": _document_preview_coordinate_space(document, "auto"),
        "highlight_strategy": "none",
        "highlight_refined": False,
        "highlight_fallback_reason": None,
        "preview_cache_object_name": str(meta.get("object_name") or ""),
    }


def _load_or_create_raw_preview_payload(
    *,
    document: dict,
    page: int,
    source_kind: str,
    cache_service: RedisCacheService,
    oss_service: MinioService,
) -> dict[str, Any]:
    if not settings.XTJS_CACHE_ENABLED:
        file_bytes, _content_type, _object_name = _load_document_source_bytes(
            document=document,
            oss_service=oss_service,
        )
        raw_payload = _raw_preview_payload_from_source(
            file_bytes=file_bytes,
            source_kind=source_kind,
            page=page,
        )
        image_bytes = bytes(raw_payload.pop("image_bytes"))
        return _raw_preview_response_payload(
            document=document,
            page=page,
            image_bytes=image_bytes,
            meta=raw_payload,
        )

    version = _preview_document_version(document)
    meta_key = cache_service.preview_meta_key(str(document.get("identifier_id")), version, page)
    cached_meta = cache_service.get_json(meta_key)
    if cached_meta.hit and isinstance(cached_meta.value, dict):
        object_name = str(cached_meta.value.get("object_name") or "").strip()
        if object_name:
            try:
                image_bytes, _content_type = oss_service.get_object_bytes(
                    object_name,
                    cached_meta.value.get("bucket_name") or None,
                )
                return _raw_preview_response_payload(
                    document=document,
                    page=page,
                    image_bytes=image_bytes,
                    meta=cached_meta.value,
                )
            except RuntimeError:
                logger.exception(
                    "preview cache object missing or unavailable document=%s page=%s object=%s",
                    document.get("identifier_id"),
                    page,
                    object_name,
                )

    file_bytes, _content_type, _object_name = _load_document_source_bytes(
        document=document,
        oss_service=oss_service,
    )
    raw_payload = _raw_preview_payload_from_source(
        file_bytes=file_bytes,
        source_kind=source_kind,
        page=page,
    )
    image_bytes = bytes(raw_payload.pop("image_bytes"))
    object_name = _preview_cache_object_name(document, version, page)
    upload_result = oss_service.upload_bytes(
        image_bytes,
        filename=f"p{int(page)}.png",
        content_type="image/png",
        object_name=object_name,
    )
    meta = {
        "object_name": upload_result["object_name"],
        "bucket_name": upload_result.get("bucket_name"),
        "page_count": raw_payload["page_count"],
        "width": raw_payload["width"],
        "height": raw_payload["height"],
        "source_kind": raw_payload["source_kind"],
        "content_type": "image/png",
        "created_at": f"{datetime.utcnow().isoformat()}Z",
    }
    cache_service.set_json(
        meta_key,
        meta,
        settings.XTJS_CACHE_PREVIEW_META_TTL_SECONDS,
    )
    return _raw_preview_response_payload(
        document=document,
        page=page,
        image_bytes=image_bytes,
        meta=meta,
    )


def _coerce_document_content(document: Optional[dict]) -> dict[str, Any]:
    raw_content = (document or {}).get("content")
    if isinstance(raw_content, dict):
        return raw_content
    if isinstance(raw_content, str):
        try:
            parsed = json.loads(raw_content)
        except (TypeError, ValueError, json.JSONDecodeError):
            return {}
        return parsed if isinstance(parsed, dict) else {}
    return {}


def _normalize_preview_coordinate_space(value: Any) -> str:
    text = str(value or "").strip().lower()
    if text in {"pdf_point", "pdf-points", "pdf_points", "pymupdf", "fitz"}:
        return "pdf_point"
    if text in {"pdf", "pdf_page", "pdf-page"}:
        return "pdf"
    if text in {"ocr", "ocr_image", "image", "image_pixel", "pixel", "pixels"}:
        return "ocr_image"
    return "auto"


def _document_preview_coordinate_space(document: Optional[dict], explicit: Any = None) -> str:
    explicit_space = _normalize_preview_coordinate_space(explicit)
    if explicit_space != "auto":
        return explicit_space
    content = _coerce_document_content(document)
    for key in ("bbox_coordinate_space", "bbox_source_coordinate_space", "coordinate_system"):
        coordinate_space = _normalize_preview_coordinate_space(content.get(key))
        if coordinate_space != "auto":
            return coordinate_space
    return "auto"


def _document_preview_kind(document: Optional[dict]) -> str:
    document_type = str((document or {}).get("document_type") or "").strip().lower()
    if document_type == DOCUMENT_TYPE_TENDER:
        return "tender"
    if document_type in BID_DOCUMENT_TYPES:
        return "bid"
    return "generic"


def _collect_preview_source_rects(
    *,
    highlight_bbox: Optional[list[float]],
    highlight_rects: Optional[list[list[float]]],
) -> list[list[float]]:
    source_rects: list[list[float]] = []
    seen_source_rects = set()
    raw_source_rects = list(highlight_rects or [])
    if highlight_bbox:
        raw_source_rects.append(highlight_bbox)
    for rect_values in raw_source_rects:
        if not isinstance(rect_values, (list, tuple)) or len(rect_values) < 4:
            continue
        try:
            key = tuple(round(float(rect_values[index]), 2) for index in range(4))
        except (TypeError, ValueError):
            continue
        if key in seen_source_rects:
            continue
        seen_source_rects.add(key)
        source_rects.append(list(key))
    return source_rects


def _render_png_data_url(image_bytes: bytes) -> str:
    return "data:image/png;base64," + base64.b64encode(image_bytes).decode("ascii")


# 将 PDF/图片的字节渲染为带 base64 的预览数据（支持高亮）
def _preview_payload_from_source(
    *,
    file_bytes: bytes,
    source_kind: str,
    page: int,
    highlight_phrases: Optional[list[str]] = None,
    highlight_bbox: Optional[list[float]] = None,
    highlight_rects: Optional[list[list[float]]] = None,
    document: Optional[dict] = None,
    highlight_coordinate_space: str = "auto",
) -> dict:
    if page <= 0:
        raise ValueError("page must be greater than 0")

    document_kind = _document_preview_kind(document)
    coordinate_space = _document_preview_coordinate_space(document, highlight_coordinate_space)
    source_rects = _collect_preview_source_rects(
        highlight_bbox=highlight_bbox,
        highlight_rects=highlight_rects,
    )
    explicit_source_rects = _collect_preview_source_rects(
        highlight_bbox=None,
        highlight_rects=highlight_rects,
    )

    if source_kind == "pdf":
        import fitz

        pdf = fitz.open(stream=file_bytes, filetype="pdf")
        try:
            page_count = int(pdf.page_count)
            if page > page_count:
                raise ValueError(f"page {page} is out of range, max page is {page_count}")
            pdf_page = pdf.load_page(page - 1)
            rect = pdf_page.rect

            direct_rects = []
            text_rects = []
            highlight_strategy = "none"
            highlight_refined = False
            highlight_fallback_reason: str | None = None
            if source_rects:
                if document_kind == "tender":
                    tender_phrases = _normalize_preview_highlight_phrases(highlight_phrases or [])
                    if tender_phrases:
                        direct_rects = _apply_pdf_text_highlights(
                            pdf_page,
                            highlight_phrases=tender_phrases,
                            highlight_bbox=source_rects[0],
                        )
                        if not direct_rects:
                            direct_rects = _apply_pdf_text_highlights(
                                pdf_page,
                                highlight_phrases=tender_phrases,
                                highlight_bbox=None,
                            )
                        if direct_rects:
                            highlight_strategy = "tender_pdf_text"
                            highlight_refined = True
                        else:
                            highlight_strategy = "tender_text_not_found"
                            highlight_fallback_reason = "tender_text_match_failed"
                    else:
                        refined_rects = source_rects
                        highlight_strategy = "tender_pdf_rect"
                        highlight_fallback_reason = None
                        direct_rects = _apply_pdf_rect_highlights(
                            pdf_page,
                            highlight_rects=refined_rects,
                            coordinate_space="pdf_point",
                        )
                elif document_kind == "bid":
                    refined_rects, highlight_fallback_reason = _refine_bid_pdf_highlights_via_fine_ocr(
                        pdf_page,
                        highlight_rects=source_rects,
                        highlight_phrases=highlight_phrases or [],
                        coordinate_space=coordinate_space,
                    )
                    highlight_refined = bool(refined_rects)
                    if refined_rects:
                        highlight_strategy = "bid_fine_ocr"
                        draw_coordinate_space = "pdf_point"
                    else:
                        refined_rects = source_rects
                        highlight_strategy = "bid_region_fallback"
                        highlight_fallback_reason = highlight_fallback_reason or "fine_ocr_no_match"
                        draw_coordinate_space = coordinate_space
                    direct_rects = _apply_pdf_rect_highlights(
                        pdf_page,
                        highlight_rects=refined_rects,
                        coordinate_space=draw_coordinate_space,
                    )
                else:
                    refined_rects = source_rects
                    highlight_strategy = "pdf_rect"
                    direct_rects = _apply_pdf_rect_highlights(
                        pdf_page,
                        highlight_rects=refined_rects,
                        coordinate_space=coordinate_space,
                    )
                if not direct_rects and source_rects:
                    highlight_fallback_reason = highlight_fallback_reason or "rects_out_of_page"

            if not direct_rects and document_kind != "bid":
                text_rects = _apply_pdf_text_highlights(
                    pdf_page,
                    highlight_phrases=highlight_phrases or [],
                    highlight_bbox=highlight_bbox,
                )
                if text_rects:
                    highlight_strategy = "pdf_text"
            elif document_kind == "bid" and not source_rects:
                highlight_strategy = "bid_no_region"
                highlight_fallback_reason = "missing_candidate_region"

            pix = pdf_page.get_pixmap(matrix=fitz.Matrix(1.8, 1.8), alpha=False)
            image_bytes = pix.tobytes("png")
            return {
                "page": page,
                "page_count": page_count,
                "width": float(rect.width),
                "height": float(rect.height),
                "image_data_url": _render_png_data_url(image_bytes),
                "source_kind": "pdf",
                "highlight_rect_count": len(text_rects) + len(direct_rects),
                "highlight_applied": bool(text_rects or direct_rects),
                "highlight_coordinate_space": coordinate_space,
                "highlight_strategy": highlight_strategy,
                "highlight_refined": highlight_refined,
                "highlight_fallback_reason": highlight_fallback_reason,
            }
        finally:
            pdf.close()

    if source_kind == "image":
        if page != 1:
            raise ValueError("image document only supports page 1 preview")
        from PIL import Image

        with Image.open(io.BytesIO(file_bytes)) as image:
            rgb = image.convert("RGB")
            direct_rects = []
            highlight_strategy = "none"
            highlight_refined = False
            highlight_fallback_reason: str | None = None
            if source_rects:
                if document_kind == "bid":
                    refined_rects, highlight_fallback_reason = _refine_bid_image_highlights_via_fine_ocr(
                        rgb,
                        highlight_rects=source_rects,
                        highlight_phrases=highlight_phrases or [],
                    )
                    highlight_refined = bool(refined_rects)
                    if refined_rects:
                        highlight_strategy = "bid_fine_ocr"
                    else:
                        refined_rects = source_rects
                        highlight_strategy = "bid_region_fallback"
                        highlight_fallback_reason = highlight_fallback_reason or "fine_ocr_no_match"
                else:
                    refined_rects = source_rects
                    highlight_strategy = "image_rect"
                direct_rects = _apply_image_rect_highlights(rgb, refined_rects)
            elif document_kind == "bid":
                highlight_strategy = "bid_no_region"
                highlight_fallback_reason = "missing_candidate_region"

            buffer = io.BytesIO()
            rgb.save(buffer, format="PNG")
            return {
                "page": 1,
                "page_count": 1,
                "width": int(rgb.width),
                "height": int(rgb.height),
                "image_data_url": _render_png_data_url(buffer.getvalue()),
                "source_kind": "image",
                "highlight_rect_count": len(direct_rects),
                "highlight_applied": bool(direct_rects),
                "highlight_coordinate_space": coordinate_space,
                "highlight_strategy": highlight_strategy,
                "highlight_refined": highlight_refined,
                "highlight_fallback_reason": highlight_fallback_reason,
            }

    raise ValueError("document source kind does not support preview")


# 归一化高亮关键词列表（去噪、去重、限制长度）
def _normalize_preview_highlight_phrases(raw_values: Any) -> list[str]:
    if raw_values is None:
        values = []
    elif isinstance(raw_values, str):
        values = [raw_values]
    else:
        values = raw_values
    normalized: list[str] = []
    for value in values or []:
        text = re.sub(r"\s+", " ", str(value or "")).strip()
        if not text:
            continue
        compact = re.sub(r"[^\w\u4e00-\u9fff]+", "", text, flags=re.UNICODE)
        if len(compact) < 2:
            continue
        if compact.isdigit():
            continue
        if text not in normalized:
            normalized.append(text[:240])
        if len(normalized) >= 12:
            break
    return normalized


# 归一化单个高亮边界框（逗号分隔的四个数字）
def _normalize_preview_highlight_bbox(raw_bbox: Any) -> Optional[list[float]]:
    if not raw_bbox:
        return None
    if isinstance(raw_bbox, (list, tuple)):
        parts = list(raw_bbox)
    else:
        parts = [segment.strip() for segment in str(raw_bbox).split(",")]
    values: list[float] = []
    for part in parts[:4]:
        try:
            values.append(float(part))
        except (TypeError, ValueError):
            return None
    if len(values) < 4:
        return None
    x0, y0, x1, y1 = values[:4]
    if x1 <= x0 or y1 <= y0:
        return None
    return [x0, y0, x1, y1]


# 生成高亮变体签名（用于缓存键，区分不同高亮参数组合）
def _preview_variant_signature(
    highlight_phrases: Optional[list[str]],
    highlight_bbox: Optional[list[float]],
    highlight_rects: Optional[list[list[float]]] = None,
    *,
    highlight_coordinate_space: str = "auto",
    document_type: str = "",
    strategy_version: str = _FINE_OCR_PREVIEW_STRATEGY_VERSION,
) -> str:
    phrases = _normalize_preview_highlight_phrases(highlight_phrases)
    bbox = highlight_bbox or []
    rects = highlight_rects or []
    coordinate_space = _normalize_preview_coordinate_space(highlight_coordinate_space)
    if not phrases and not bbox and not rects:
        return ""
    payload = json.dumps(
        {
            "phrases": phrases,
            "bbox": [round(float(value), 2) for value in bbox],
            "rects": [
                [round(float(value), 2) for value in rect[:4]]
                for rect in rects
                if isinstance(rect, (list, tuple)) and len(rect) >= 4
            ],
            "coordinate_space": coordinate_space,
            "document_type": str(document_type or "").strip().lower(),
            "strategy_version": strategy_version,
        },
        ensure_ascii=False,
        sort_keys=True,
    )
    return hashlib.sha1(payload.encode("utf-8")).hexdigest()


# 从短语中提取用于匹配的 token（去标点、小写、排序）
def _highlight_tokens_from_phrases(phrases: list[str]) -> list[str]:
    tokens: list[str] = []
    for phrase in phrases:
        for raw_token in _HIGHLIGHT_TOKEN_PATTERN.findall(str(phrase or "")):
            token = re.sub(r"[^\w\u4e00-\u9fff]+", "", raw_token, flags=re.UNICODE).lower()
            if len(token) < 2:
                continue
            if token.isdigit():
                continue
            if token not in tokens:
                tokens.append(token)
    tokens.sort(key=len, reverse=True)  # 长 token 优先匹配
    return tokens[:48]


# 判断 PDF 中的单词是否匹配任一高亮 token
def _word_matches_highlight_token(word_text: str, tokens: list[str]) -> bool:
    normalized = re.sub(r"[^\w\u4e00-\u9fff]+", "", str(word_text or ""), flags=re.UNICODE).lower()
    if len(normalized) < 2:
        return False
    for token in tokens:
        if normalized == token:
            return True
        if len(token) >= 4 and normalized in token:
            return True
        if len(normalized) >= 4 and token in normalized:
            return True
    return False


# 收集 PDF 页内需要高亮的矩形（基于文本词级定位）
def _normalized_pdf_highlight_text(text: Any) -> str:
    return re.sub(r"[^\w\u4e00-\u9fff]+", "", str(text or ""), flags=re.UNICODE).lower()


def _pdf_dict_line_text(line: dict) -> str:
    spans = line.get("spans") or []
    return "".join(str(span.get("text") or "") for span in spans if isinstance(span, dict))


def _pdf_dict_line_rect(line: dict):
    import fitz

    bbox = line.get("bbox")
    if isinstance(bbox, (list, tuple)) and len(bbox) >= 4:
        try:
            return fitz.Rect([float(value) for value in bbox[:4]])
        except (TypeError, ValueError):
            pass

    rect = None
    for span in line.get("spans") or []:
        if not isinstance(span, dict):
            continue
        span_bbox = span.get("bbox")
        if not isinstance(span_bbox, (list, tuple)) or len(span_bbox) < 4:
            continue
        try:
            span_rect = fitz.Rect([float(value) for value in span_bbox[:4]])
        except (TypeError, ValueError):
            continue
        rect = span_rect if rect is None else rect | span_rect
    return rect


def _collect_pdf_highlight_rects(
    pdf_page,
    *,
    highlight_phrases: list[str],
    highlight_bbox: Optional[list[float]],
):
    import fitz

    tokens = _highlight_tokens_from_phrases(highlight_phrases)
    if not tokens:
        return []

    coerced_bbox = _coerce_rect_to_pdf_page_space(pdf_page, highlight_bbox) if highlight_bbox else None
    bbox_rect = fitz.Rect(coerced_bbox) if coerced_bbox else None
    phrase_keys = [
        key
        for phrase in highlight_phrases or []
        if (key := _normalized_pdf_highlight_text(phrase))
    ]

    def line_matches(line_text: str, *, restrict_bbox: bool) -> bool:
        line_key = _normalized_pdf_highlight_text(line_text)
        if len(line_key) < 2:
            return False
        if any(phrase_key in line_key for phrase_key in phrase_keys):
            return True
        if restrict_bbox:
            return any(
                token in line_key or (len(line_key) >= 4 and line_key in token)
                for token in tokens
            )
        return False

    def iter_line_matches(restrict_bbox: bool):
        try:
            text_dict = pdf_page.get_text("dict", sort=True) or {}
        except TypeError:
            text_dict = pdf_page.get_text("dict") or {}
        matched = []
        for block in text_dict.get("blocks") or []:
            if not isinstance(block, dict) or int(block.get("type") or 0) != 0:
                continue
            for line in block.get("lines") or []:
                if not isinstance(line, dict):
                    continue
                rect = _pdf_dict_line_rect(line)
                if rect is None or rect.is_empty:
                    continue
                if restrict_bbox and bbox_rect is not None:
                    overlap = rect & bbox_rect
                    if overlap.is_empty or overlap.get_area() <= 0:
                        continue
                if not line_matches(_pdf_dict_line_text(line), restrict_bbox=restrict_bbox):
                    continue
                matched.append(rect)
        return matched

    line_rects = iter_line_matches(restrict_bbox=True)
    if not line_rects and bbox_rect is not None:
        line_rects = iter_line_matches(restrict_bbox=False)
    if line_rects:
        deduped_line_rects = []
        seen_line_rects = set()
        for rect in line_rects:
            key = tuple(round(float(value), 2) for value in (rect.x0, rect.y0, rect.x1, rect.y1))
            if key in seen_line_rects:
                continue
            seen_line_rects.add(key)
            deduped_line_rects.append(rect)
        return deduped_line_rects

    def iter_matches(restrict_bbox: bool):
        matched = []
        for word in pdf_page.get_text("words", sort=True):
            if len(word) < 8:
                continue
            rect = fitz.Rect(word[:4])
            if restrict_bbox and bbox_rect is not None:
                overlap = rect & bbox_rect
                if overlap.is_empty:
                    continue
                if overlap.get_area() <= 0:
                    continue
            if not _word_matches_highlight_token(word[4], tokens):
                continue
            matched.append(
                {
                    "rect": rect,
                    "block": int(word[5]),
                    "line": int(word[6]),
                    "word": int(word[7]),
                }
            )
        return matched

    # 先在指定 bbox 内匹配，若无效再全文匹配
    matches = iter_matches(restrict_bbox=True)
    if not matches and bbox_rect is not None:
        matches = iter_matches(restrict_bbox=False)

    if not matches:
        return []

    # 合并同行相邻单词的矩形
    merged: list[fitz.Rect] = []
    current_rect = None
    current_key = None
    current_word_index = None

    for item in matches:
        rect = item["rect"]
        key = (item["block"], item["line"])
        word_index = item["word"]
        if current_rect is None:
            current_rect = fitz.Rect(rect)
            current_key = key
            current_word_index = word_index
            continue
        gap = rect.x0 - current_rect.x1
        same_line = key == current_key
        if same_line and word_index <= (current_word_index or 0) + 2 and gap <= max(18.0, rect.height * 1.2):
            current_rect |= rect
            current_word_index = word_index
            continue
        merged.append(current_rect)
        current_rect = fitz.Rect(rect)
        current_key = key
        current_word_index = word_index

    if current_rect is not None:
        merged.append(current_rect)
    return merged


# 在 PDF 页面上绘制文本匹配的高亮矩形
def _apply_pdf_text_highlights(pdf_page, *, highlight_phrases: list[str], highlight_bbox: Optional[list[float]]):
    rects = _collect_pdf_highlight_rects(
        pdf_page,
        highlight_phrases=_normalize_preview_highlight_phrases(highlight_phrases),
        highlight_bbox=highlight_bbox,
    )
    if not rects:
        return []
    for rect in rects:
        pdf_page.draw_rect(
            rect,
            color=(1.0, 0.91, 0.2),
            fill=(1.0, 0.91, 0.2),
            width=0.3,
            fill_opacity=0.28,
            overlay=True,
        )
    return rects


# 归一化 JSON 格式的高亮矩形数组
def _normalize_preview_highlight_rects(raw_value: Any) -> list[list[float]]:
    if not raw_value:
        return []
    if isinstance(raw_value, (list, tuple)):
        parsed = raw_value
    else:
        try:
            parsed = json.loads(str(raw_value))
        except (TypeError, ValueError, json.JSONDecodeError):
            return []
    rects: list[list[float]] = []
    for item in parsed if isinstance(parsed, list) else []:
        if not isinstance(item, (list, tuple)) or len(item) < 4:
            continue
        values = []
        try:
            values = [float(item[index]) for index in range(4)]
        except (TypeError, ValueError):
            continue
        x0, y0, x1, y1 = values
        if x1 <= x0 or y1 <= y0:
            continue
        rects.append([x0, y0, x1, y1])
    return rects[:24]


# 将外部传来的矩形坐标转换到 PDF 页面坐标系（处理缩放）
def _coerce_rect_to_pdf_page_space(
    pdf_page,
    rect_values: list[float],
    coordinate_space: str = "auto",
) -> Optional[list[float]]:
    import fitz

    if not isinstance(rect_values, (list, tuple)) or len(rect_values) < 4:
        return None
    try:
        x0, y0, x1, y1 = [float(rect_values[index]) for index in range(4)]
    except (TypeError, ValueError):
        return None
    if x1 <= x0 or y1 <= y0:
        return None

    page_rect = fitz.Rect(pdf_page.rect)
    page_width = max(float(page_rect.width), 1.0)
    page_height = max(float(page_rect.height), 1.0)

    normalized_coordinate_space = _normalize_preview_coordinate_space(coordinate_space)

    # pdf_point/pdf 已经是 PyMuPDF 页面坐标，避免自动缩放造成二次偏移。
    scale = 1.0
    if normalized_coordinate_space not in {"pdf_point", "pdf"}:
        max_ratio = max(x1 / page_width, y1 / page_height)
        if max_ratio > 1.05:
            for candidate in (1.5, 2.0, 3.0, 4.0):
                if (x1 / candidate) <= page_width * 1.05 and (y1 / candidate) <= page_height * 1.05:
                    scale = candidate
                    break
            else:
                scale = max_ratio

    x0 /= scale
    y0 /= scale
    x1 /= scale
    y1 /= scale

    # 裁剪到页面范围内
    x0 = min(max(x0, 0.0), page_width)
    y0 = min(max(y0, 0.0), page_height)
    x1 = min(max(x1, 0.0), page_width)
    y1 = min(max(y1, 0.0), page_height)

    if x1 <= x0 or y1 <= y0:
        return None
    return [x0, y0, x1, y1]


def _coerce_rect_to_image_space(image, rect_values: list[float]) -> Optional[list[float]]:
    if not isinstance(rect_values, (list, tuple)) or len(rect_values) < 4:
        return None
    try:
        x0, y0, x1, y1 = [float(rect_values[index]) for index in range(4)]
    except (TypeError, ValueError):
        return None
    if x1 <= x0 or y1 <= y0:
        return None
    width = max(float(getattr(image, "width", 0) or 0), 1.0)
    height = max(float(getattr(image, "height", 0) or 0), 1.0)
    x0 = min(max(x0, 0.0), width)
    y0 = min(max(y0, 0.0), height)
    x1 = min(max(x1, 0.0), width)
    y1 = min(max(y1, 0.0), height)
    if x1 <= x0 or y1 <= y0:
        return None
    return [x0, y0, x1, y1]


def _expand_pdf_clip(pdf_page, rect_values: list[float], padding: float = 2.0):
    import fitz

    rect = fitz.Rect(rect_values)
    clip = fitz.Rect(rect)
    clip.x0 = max(0, clip.x0 - padding)
    clip.y0 = max(0, clip.y0 - padding)
    clip.x1 = min(float(pdf_page.rect.width), clip.x1 + padding)
    clip.y1 = min(float(pdf_page.rect.height), clip.y1 + padding)
    return clip if clip.x1 > clip.x0 and clip.y1 > clip.y0 else None


def _expand_image_clip(image, rect_values: list[float], padding: float = 4.0) -> Optional[tuple[int, int, int, int]]:
    x0, y0, x1, y1 = rect_values[:4]
    width = int(getattr(image, "width", 0) or 0)
    height = int(getattr(image, "height", 0) or 0)
    left = max(0, int(round(float(x0) - padding)))
    top = max(0, int(round(float(y0) - padding)))
    right = min(width, int(round(float(x1) + padding)))
    bottom = min(height, int(round(float(y1) + padding)))
    if right <= left or bottom <= top:
        return None
    return left, top, right, bottom


def _map_fine_ocr_boxes_to_pdf_rects(
    *,
    items: list[dict[str, Any]],
    clip,
    scale: float,
    highlight_phrases: list[str],
) -> list[list[float]]:
    rects: list[list[float]] = []
    for item in items or []:
        section_text = str(item.get("text") or item.get("raw_text") or "").strip()
        bbox = item.get("bbox")
        if not section_text or not isinstance(bbox, (list, tuple)) or len(bbox) < 4:
            continue
        local_bbox = [float(bbox[index]) for index in range(4)]
        for local_rect in _section_subrects_for_phrases(section_text, local_bbox, highlight_phrases):
            rects.append(
                [
                    clip.x0 + (local_rect[0] / scale),
                    clip.y0 + (local_rect[1] / scale),
                    clip.x0 + (local_rect[2] / scale),
                    clip.y0 + (local_rect[3] / scale),
                ]
            )
    return rects


def _map_fine_ocr_boxes_to_image_rects(
    *,
    items: list[dict[str, Any]],
    clip: tuple[int, int, int, int],
    highlight_phrases: list[str],
) -> list[list[float]]:
    rects: list[list[float]] = []
    left, top, _right, _bottom = clip
    for item in items or []:
        section_text = str(item.get("text") or item.get("raw_text") or "").strip()
        bbox = item.get("bbox")
        if not section_text or not isinstance(bbox, (list, tuple)) or len(bbox) < 4:
            continue
        local_bbox = [float(bbox[index]) for index in range(4)]
        for local_rect in _section_subrects_for_phrases(section_text, local_bbox, highlight_phrases):
            rects.append(
                [
                    left + local_rect[0],
                    top + local_rect[1],
                    left + local_rect[2],
                    top + local_rect[3],
                ]
            )
    return rects


def _run_fine_text_ocr_on_png(png_bytes: bytes) -> tuple[list[dict[str, Any]], str | None]:
    ocr_service = _get_preview_ocr_service()
    if ocr_service is None:
        return [], "fine_ocr_unavailable"
    if not hasattr(ocr_service, "extract_text_boxes_without_layout"):
        return [], "fine_ocr_method_unavailable"

    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_file:
        temp_file.write(png_bytes)
        temp_path = temp_file.name
    try:
        payload = ocr_service.extract_text_boxes_without_layout(temp_path)
    except Exception as exc:
        return [], f"fine_ocr_failed:{exc}"
    finally:
        try:
            os.unlink(temp_path)
        except OSError:
            pass

    if not isinstance(payload, dict):
        return [], "fine_ocr_empty_payload"
    return [item for item in payload.get("items") or [] if isinstance(item, dict)], None


def _refine_bid_pdf_highlights_via_fine_ocr(
    pdf_page,
    *,
    highlight_rects: list[list[float]],
    highlight_phrases: list[str],
    coordinate_space: str,
) -> tuple[list[list[float]], str | None]:
    import fitz

    phrases = _normalize_preview_highlight_phrases(highlight_phrases)
    if not phrases:
        return [], "missing_highlight_phrase"
    if not highlight_rects:
        return [], "missing_candidate_region"

    scale = 2.0
    refined: list[list[float]] = []
    fallback_reason: str | None = None
    for rect_values in highlight_rects[:4]:
        coerced = _coerce_rect_to_pdf_page_space(pdf_page, rect_values, coordinate_space)
        if not coerced:
            fallback_reason = "candidate_region_out_of_page"
            continue
        clip = _expand_pdf_clip(pdf_page, coerced)
        if clip is None:
            fallback_reason = "candidate_region_out_of_page"
            continue
        try:
            pix = pdf_page.get_pixmap(matrix=fitz.Matrix(scale, scale), clip=clip, alpha=False)
            pix_bytes = pix.tobytes("png")
        except Exception as exc:
            fallback_reason = f"clip_render_failed:{exc}"
            continue

        items, reason = _run_fine_text_ocr_on_png(pix_bytes)
        if reason:
            fallback_reason = reason
            continue
        refined.extend(
            _map_fine_ocr_boxes_to_pdf_rects(
                items=items,
                clip=clip,
                scale=scale,
                highlight_phrases=phrases,
            )
        )

    return refined, None if refined else (fallback_reason or "fine_ocr_no_match")


def _refine_bid_image_highlights_via_fine_ocr(
    image,
    *,
    highlight_rects: list[list[float]],
    highlight_phrases: list[str],
) -> tuple[list[list[float]], str | None]:
    phrases = _normalize_preview_highlight_phrases(highlight_phrases)
    if not phrases:
        return [], "missing_highlight_phrase"
    if not highlight_rects:
        return [], "missing_candidate_region"

    refined: list[list[float]] = []
    fallback_reason: str | None = None
    for rect_values in highlight_rects[:4]:
        coerced = _coerce_rect_to_image_space(image, rect_values)
        if not coerced:
            fallback_reason = "candidate_region_out_of_page"
            continue
        clip = _expand_image_clip(image, coerced)
        if clip is None:
            fallback_reason = "candidate_region_out_of_page"
            continue
        try:
            crop = image.crop(clip)
            buffer = io.BytesIO()
            crop.save(buffer, format="PNG")
            pix_bytes = buffer.getvalue()
        except Exception as exc:
            fallback_reason = f"clip_render_failed:{exc}"
            continue

        items, reason = _run_fine_text_ocr_on_png(pix_bytes)
        if reason:
            fallback_reason = reason
            continue
        refined.extend(
            _map_fine_ocr_boxes_to_image_rects(
                items=items,
                clip=clip,
                highlight_phrases=phrases,
            )
        )

    return refined, None if refined else (fallback_reason or "fine_ocr_no_match")


def _apply_image_rect_highlights(image, rects: list[list[float]]) -> list[list[float]]:
    from PIL import Image, ImageDraw

    applied = []
    if not rects:
        return applied
    overlay = Image.new("RGBA", image.size, (0, 0, 0, 0))
    draw = ImageDraw.Draw(overlay)
    for rect_values in rects:
        coerced = _coerce_rect_to_image_space(image, rect_values)
        if not coerced:
            continue
        applied.append(coerced)
        draw.rectangle(coerced, fill=(255, 237, 77, 70), outline=(245, 158, 11, 180), width=2)
    if applied:
        composite = Image.alpha_composite(image.convert("RGBA"), overlay)
        if image.mode == "RGBA":
            image.paste(composite)
        else:
            image.paste(composite.convert("RGB"))
    return applied


# 在 PDF 页面上绘制用户指定的矩形高亮
def _apply_pdf_rect_highlights(
    pdf_page,
    *,
    highlight_rects: list[list[float]],
    coordinate_space: str = "auto",
):
    import fitz

    rects = []
    for item in highlight_rects or []:
        coerced = _coerce_rect_to_pdf_page_space(pdf_page, item, coordinate_space)
        if not coerced:
            continue
        rect = fitz.Rect(coerced)
        if rect.is_empty or rect.get_area() <= 0:
            continue
        pdf_page.draw_rect(
            rect,
            color=(1.0, 0.93, 0.3),
            fill=(1.0, 0.93, 0.3),
            width=0.2,
            fill_opacity=0.22,
            overlay=True,
        )
        rects.append(rect)
    return rects


# 获取可用的 OCR 服务实例（用于预览时的 OCR 辅助定位）
def _get_preview_ocr_service():
    try:
        analysis_service = get_text_analysis_service()
    except Exception:
        return None

    direct = getattr(analysis_service, "ocr_service", None)
    if direct is not None and bool(getattr(direct, "available", False)):
        return direct

    services = getattr(analysis_service, "_services", None)
    if isinstance(services, list):
        for service in services:
            ocr_service = getattr(service, "ocr_service", None)
            if ocr_service is not None and bool(getattr(ocr_service, "available", False)):
                return ocr_service
    return None


# 紧凑化文本（去除所有非字母数字汉字的字符）
def _compact_highlight_text(text: str) -> str:
    return re.sub(r"[^\w\u4e00-\u9fff]+", "", str(text or ""), flags=re.UNICODE).lower()


# 在文本中查找短语的紧凑匹配范围（用于 OCR 结果中定位高亮短语）
def _find_compact_substring_ranges(text: str, phrase: str) -> list[tuple[int, int]]:
    source = str(text or "")
    target = str(phrase or "")
    compact_source_chars: list[str] = []
    compact_to_source_index: list[int] = []
    for index, char in enumerate(source):
        if re.match(r"[\w\u4e00-\u9fff]", char, flags=re.UNICODE):
            compact_source_chars.append(char.lower())
            compact_to_source_index.append(index)
    compact_source = "".join(compact_source_chars)
    compact_target = _compact_highlight_text(target)
    if not compact_source or not compact_target:
        return []

    ranges: list[tuple[int, int]] = []
    search_from = 0
    while search_from < len(compact_source):
        found_at = compact_source.find(compact_target, search_from)
        if found_at < 0:
            break
        end_at = found_at + len(compact_target)
        ranges.append((found_at, end_at))
        search_from = max(found_at + 1, end_at)
    return ranges


# 在 OCR 区段内按词语匹配生成子矩形（用于 OCR 精修高亮）
def _section_subrects_for_phrases(section_text: str, section_bbox: list[float], phrases: list[str]) -> list[list[float]]:
    if not section_text or not phrases:
        return []
    x0, y0, x1, y1 = section_bbox[:4]
    width = max(float(x1) - float(x0), 1.0)
    compact_length = len(_compact_highlight_text(section_text))
    if compact_length <= 0:
        return []

    rects: list[list[float]] = []
    seen = set()
    for phrase in phrases:
        for start_at, end_at in _find_compact_substring_ranges(section_text, phrase):
            left = x0 + width * (start_at / compact_length)
            right = x0 + width * (end_at / compact_length)
            key = (round(left, 1), round(right, 1), round(y0, 1), round(y1, 1))
            if key in seen:
                continue
            seen.add(key)
            rects.append([left, y0, max(left + 2.0, right), y1])
    return rects


# 当文本高亮失败时，利用 OCR 对 PDF 页面截图进行识别，并精修高亮区域
def _refine_highlight_rects_via_ocr(pdf_page, *, highlight_rects: list[list[float]], highlight_phrases: list[str]) -> list[list[float]]:
    ocr_service = _get_preview_ocr_service()
    if ocr_service is None:
        return highlight_rects

    import fitz

    scale = 2.0
    refined: list[list[float]] = []
    phrases = _normalize_preview_highlight_phrases(highlight_phrases)
    if not phrases:
        return highlight_rects
    if len(highlight_rects or []) > 1:
        return highlight_rects

    for rect_values in highlight_rects or []:
        coerced = _coerce_rect_to_pdf_page_space(pdf_page, rect_values)
        if not coerced:
            continue
        rect = fitz.Rect(coerced)
        if rect.is_empty or rect.get_area() <= 0:
            continue
        clip = fitz.Rect(rect)
        clip.x0 = max(0, clip.x0 - 2)
        clip.y0 = max(0, clip.y0 - 2)
        clip.x1 = min(float(pdf_page.rect.width), clip.x1 + 2)
        clip.y1 = min(float(pdf_page.rect.height), clip.y1 + 2)
        if clip.x1 <= clip.x0 or clip.y1 <= clip.y0:
            refined.append(coerced)
            continue
        try:
            pix = pdf_page.get_pixmap(matrix=fitz.Matrix(scale, scale), clip=clip, alpha=False)
            pix_bytes = pix.tobytes("png")
        except Exception:
            refined.append(coerced)
            continue
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_file:
            temp_file.write(pix_bytes)
            temp_path = temp_file.name
        try:
            ocr_payload = ocr_service.extract_all(temp_path, "png")
        except Exception:
            ocr_payload = {}
        finally:
            try:
                os.unlink(temp_path)
            except OSError:
                pass

        sections = []
        if isinstance(ocr_payload, dict):
            sections = list(ocr_payload.get("layout_sections") or [])
        matched_any = False
        for section in sections:
            section_text = str(section.get("text") or section.get("raw_text") or "").strip()
            bbox = section.get("bbox") or section.get("bbox_ocr") or section.get("box")
            if not section_text or not isinstance(bbox, (list, tuple)) or len(bbox) < 4:
                continue
            local_bbox = [float(bbox[index]) for index in range(4)]
            rects = _section_subrects_for_phrases(section_text, local_bbox, phrases)
            for local_rect in rects:
                refined.append(
                    [
                        clip.x0 + (local_rect[0] / scale),
                        clip.y0 + (local_rect[1] / scale),
                        clip.x0 + (local_rect[2] / scale),
                        clip.y0 + (local_rect[3] / scale),
                    ]
                )
                matched_any = True
        if not matched_any:
            refined.append([clip.x0, clip.y0, clip.x1, clip.y1])

    return refined or highlight_rects


# 执行项目重复检查，并持久化结果和合并聚类
def _run_project_duplicate_check(
    *,
    identifier_id: str,
    document_types: Optional[list[str]],
    max_evidence_sections: int,
    max_pairs_per_type: int,
    result_key: str,
    db_service: PostgreSQLService,
    duplicate_check_service: DuplicateCheckService,
    duplicate_scope: Optional[str] = None,
    persist_to_latest: bool = False,
):
    payload_data = ProjectAnalysisInputLoader(db_service).load(identifier_id)
    if not payload_data:
        raise HTTPException(status_code=404, detail="项目不存在")

    duplicate_result = duplicate_check_service.check_project_documents(
        project_identifier=identifier_id,
        project=payload_data["project"],
        document_records=payload_data["documents"],
        document_types=document_types,
        max_evidence_sections=max_evidence_sections,
        max_pairs_per_type=max_pairs_per_type,
        duplicate_scope=duplicate_scope,
    )
    if persist_to_latest:
        db_service.update_project_manual_review_result(
            project_identifier_id=identifier_id,
            result_key=result_key,
            result_value=duplicate_result,
        )
    else:
        db_service.upsert_project_result_item(
            project_identifier_id=identifier_id,
            result_key=result_key,
            result_value=duplicate_result,
        )
        _persist_duplicate_merge_results(
            db_service=db_service,
            project_identifier=identifier_id,
            source_result_key=result_key,
            raw_result=duplicate_result,
        )
    _invalidate_project_cache_by_identifier(identifier_id)
    return duplicate_result


def _run_project_deviation_check(
    *,
    identifier_id: str,
    db_service: PostgreSQLService,
    persist_to_latest: bool = False,
) -> dict:
    review_service = UnifiedBusinessReviewService(db_service=db_service)
    payload_data = _load_project_review_documents(
        identifier_id=identifier_id,
        db_service=db_service,
        include_excluded=False,
    )
    project = db_service.get_project_by_identifier(identifier_id)
    if not project:
        raise ValueError(f"project not found: {identifier_id}")
    review = review_service._review_project_deviation_documents(
        project_identifier=identifier_id,
        payload_data=payload_data,
    )
    if persist_to_latest:
        result_record = db_service.update_project_manual_review_result(
            project_identifier_id=identifier_id,
            result_key="deviation_check",
            result_value=review,
        )
    else:
        result_record = db_service.upsert_project_result_item(
            identifier_id,
            "deviation_check",
            review,
        )
    _invalidate_project_cache_by_identifier(identifier_id)
    return {
        "project": project,
        "result_key": "deviation_check",
        "overview": review_service._build_response_overview(review),
        "review": review,
        "result_record": result_record,
    }


def _load_project_review_documents(
    *,
    identifier_id: str,
    db_service: PostgreSQLService,
    include_excluded: bool = False,
) -> dict[str, Any]:
    payload_data = ProjectAnalysisInputLoader(db_service).load(
        identifier_id,
        include_excluded=include_excluded,
    )
    if not payload_data:
        raise HTTPException(status_code=404, detail="项目不存在")
    return payload_data



# 项目人员复用检查
def _run_project_personnel_reuse_check(
    *,
    identifier_id: str,
    db_service: PostgreSQLService,
    bid_document_review_service: BidDocumentReviewService,
    confirmed_names: Optional[list[Any]] = None,
    persist_to_latest: bool = False,
) -> dict:
    payload_data = _load_project_review_documents(
        identifier_id=identifier_id,
        db_service=db_service,
    )
    personnel_result = bid_document_review_service.check_project_personnel_reuse(
        project_identifier=identifier_id,
        project=payload_data["project"],
        document_records=payload_data["documents"],
        document_types=None,
        confirmed_names=confirmed_names,
    )
    if persist_to_latest:
        db_service.update_project_manual_review_result(
            project_identifier_id=identifier_id,
            result_key="personnel_reuse_check",
            result_value=personnel_result,
        )
    else:
        db_service.upsert_project_result_item(
            project_identifier_id=identifier_id,
            result_key="personnel_reuse_check",
            result_value=personnel_result,
        )
    if confirmed_names is not None:
        db_service.update_project_manual_review_result(
            project_identifier_id=identifier_id,
            result_key="personnel_reuse_check",
            result_value=personnel_result,
        )
        for item in confirmed_names:
            if isinstance(item, dict):
                document_id = str(item.get("document_identifier_id") or item.get("identifier_id") or "").strip()
                if document_id:
                    DocumentWorkingCopyService(db_service).apply_personnel_reuse_check(
                        document_id,
                        {
                            "schema_version": "1.0",
                            "confirmation_status": "confirmed",
                            "confirmed_names": confirmed_names,
                        },
                    )
                    break
    _invalidate_project_cache_by_identifier(identifier_id)
    return personnel_result


def _normalize_personnel_draft_documents(
    draft_documents: list[dict[str, Any]],
    project_documents: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    metadata_by_id: dict[str, dict[str, Any]] = {}
    for document in project_documents:
        if not isinstance(document, dict):
            continue
        document_id = str(document.get("identifier_id") or document.get("document_id") or "").strip()
        if not document_id:
            continue
        metadata = {
            "identifier_id": document_id,
            "document_type": str(document.get("relation_role") or document.get("document_type") or "").strip(),
            "file_name": str(document.get("file_name") or "").strip(),
            "relation_id": document.get("relation_id"),
        }
        metadata_by_id[document_id] = metadata

    def apply_metadata(document: dict[str, Any]) -> dict[str, Any]:
        normalized = deepcopy(document)
        document_id = str(
            normalized.get("document_identifier_id")
            or normalized.get("identifier_id")
            or normalized.get("doc_id")
            or ""
        ).strip()
        metadata = metadata_by_id.get(document_id)
        if not metadata:
            return normalized
        normalized["document_identifier_id"] = metadata["identifier_id"]
        normalized["identifier_id"] = str(normalized.get("identifier_id") or metadata["identifier_id"])
        normalized["document_type"] = metadata["document_type"] or normalized.get("document_type")
        if metadata.get("file_name") and not str(normalized.get("file_name") or "").strip():
            normalized["file_name"] = metadata["file_name"]
        if metadata.get("relation_id") is not None and normalized.get("relation_id") in (None, ""):
            normalized["relation_id"] = metadata["relation_id"]

        entries = normalized.get("personnel_entries")
        if isinstance(entries, list):
            normalized["personnel_entries"] = [
                _normalize_personnel_draft_entry_metadata(entry, normalized)
                if isinstance(entry, dict)
                else entry
                for entry in entries
            ]
        return normalized

    return [
        apply_metadata(document) if isinstance(document, dict) else document
        for document in draft_documents
    ]


def _normalize_personnel_draft_entry_metadata(
    entry: dict[str, Any],
    document: dict[str, Any],
) -> dict[str, Any]:
    normalized = deepcopy(entry)
    document_id = str(document.get("document_identifier_id") or document.get("identifier_id") or "").strip()
    if document_id:
        normalized["document_identifier_id"] = str(normalized.get("document_identifier_id") or document_id)
    if document.get("document_type"):
        normalized["document_type"] = document.get("document_type")
    if document.get("file_name") and not str(normalized.get("file_name") or "").strip():
        normalized["file_name"] = document.get("file_name")
    if document.get("relation_id") is not None and normalized.get("relation_id") in (None, ""):
        normalized["relation_id"] = document.get("relation_id")
    return normalized


def _persist_project_personnel_reuse_draft(
    *,
    identifier_id: str,
    db_service: PostgreSQLService,
    bid_document_review_service: BidDocumentReviewService,
    draft_documents: list[dict[str, Any]],
    confirmation_status: str,
) -> dict:
    payload_data = _load_project_review_documents(
        identifier_id=identifier_id,
        db_service=db_service,
    )
    draft_documents = _normalize_personnel_draft_documents(
        draft_documents,
        payload_data["documents"],
    )
    personnel_result = bid_document_review_service.build_personnel_reuse_from_draft(
        project_identifier=identifier_id,
        project=payload_data["project"],
        documents=draft_documents,
        confirmation_status=confirmation_status,
    )
    for document in draft_documents:
        if not isinstance(document, dict):
            continue
        document_id = str(document.get("document_identifier_id") or document.get("identifier_id") or "").strip()
        if not document_id:
            continue
        DocumentWorkingCopyService(db_service).apply_personnel_reuse_check(
            document_id,
            {
                "schema_version": "1.0",
                "confirmation_status": confirmation_status,
                "document": document,
                "documents": [document],
            },
        )
    if confirmation_status == "confirmed":
        db_service.update_project_manual_review_result(
            project_identifier_id=identifier_id,
            result_key="personnel_reuse_check",
            result_value=personnel_result,
        )
    _invalidate_project_cache_by_identifier(identifier_id)
    return personnel_result


# 构建项目快照（仅含标识）
def _project_snapshot(project_identifier: str) -> dict:
    return {"identifier_id": project_identifier}


# 为项目生成相关 API 链接
def _project_api_links(project_identifier: str) -> dict[str, str]:
    quoted_identifier = quote(project_identifier, safe="")
    return {
        "detail_url": f"/api/postgresql/projects/{quoted_identifier}",
        "results_url": f"/api/postgresql/projects/{quoted_identifier}/results",
        "merged_results_url": f"/api/postgresql/projects/{quoted_identifier}/merged-results",
        "visualization_url": f"/api/postgresql/projects/{quoted_identifier}/visualization-data",
    }


# 为文档生成相关 API 链接
def _document_api_links(document_identifier: str) -> dict[str, str]:
    quoted_identifier = quote(document_identifier, safe="")
    return {
        "detail_url": f"/api/postgresql/documents/{quoted_identifier}",
        "source_url": f"/api/postgresql/documents/{quoted_identifier}/source",
        "preview_url_template": f"/api/postgresql/documents/{quoted_identifier}/preview/pages/{{page}}",
    }


# 从项目详情的关系中收集所有文档标识
def _collect_project_document_identifiers(project_detail: dict[str, Any]) -> list[str]:
    identifiers: list[str] = []
    seen: set[str] = set()
    for relation in project_detail.get("relations") or []:
        for field_name in (
            "tender_identifier_id",
            "business_bid_identifier_id",
            "technical_bid_identifier_id",
        ):
            identifier = str(relation.get(field_name) or "").strip()
            if not identifier or identifier in seen:
                continue
            seen.add(identifier)
            identifiers.append(identifier)
    return identifiers


# 提取结果记录的元信息（轻量版）
def _build_result_record_meta(result_record: Optional[dict[str, Any]]) -> Optional[dict[str, Any]]:
    if not isinstance(result_record, dict) or not result_record:
        return None
    return {
        "id": result_record.get("id"),
        "project_identifier_id": result_record.get("project_identifier_id"),
        "create_time": result_record.get("create_time"),
        "update_time": result_record.get("update_time"),
    }


# 压缩显示结果（移除原始的查重合并键）
def _compact_display_results_for_response(display_results: dict[str, Any]) -> dict[str, Any]:
    compact_results = dict(display_results)
    compact_results.pop("duplicate_check", None)
    for merged_key in MERGED_RESULT_KEY_BY_DOC_TYPE.values():
        compact_results.pop(merged_key, None)
    return _filter_visible_result_keys(compact_results)


def _parse_display_result_path(path: str) -> list[Any]:
    parts: list[Any] = []
    token = ""
    i = 0
    while i < len(path):
        ch = path[i]
        if ch == ".":
            if token:
                parts.append(token)
                token = ""
            i += 1
            continue
        if ch == "[":
            if token:
                parts.append(token)
                token = ""
            end = path.find("]", i)
            if end < 0:
                return []
            index_text = path[i + 1:end].strip()
            if not index_text.isdigit():
                return []
            parts.append(int(index_text))
            i = end + 1
            continue
        token += ch
        i += 1
    if token:
        parts.append(token)
    return parts


def _set_display_result_path_value(root: Any, path: str, value: Any) -> bool:
    parts = _parse_display_result_path(path)
    if not parts:
        return False
    node = root
    for part in parts[:-1]:
        if isinstance(part, int):
            if not isinstance(node, list) or part < 0 or part >= len(node):
                return False
            node = node[part]
            continue
        if not isinstance(node, dict) or part not in node:
            return False
        node = node[part]
    last = parts[-1]
    if isinstance(last, int):
        if not isinstance(node, list) or last < 0 or last >= len(node):
            return False
        node[last] = value
        return True
    if not isinstance(node, dict):
        return False
    node[last] = value
    return True


def _apply_manual_result_inputs(payload: Any, manual_payload: Any) -> Any:
    if not isinstance(payload, dict) or not isinstance(manual_payload, dict):
        return payload
    corrected = deepcopy(payload)
    applied_items: list[dict[str, Any]] = []
    for item in manual_payload.get("items") or []:
        if not isinstance(item, dict):
            continue
        if "manual_value" not in item or item.get("manual_value") is None:
            continue
        result_path = str(item.get("result_path") or "").strip()
        if not result_path:
            continue
        if _set_display_result_path_value(corrected, result_path, item.get("manual_value")):
            applied_items.append(item)
    if applied_items:
        corrected["manual_review_summary"] = {
            "applied": True,
            "applied_item_count": len(applied_items),
            "stored_item_count": len([item for item in manual_payload.get("items") or [] if isinstance(item, dict)]),
        }
        corrected["review_mode"] = "manual_corrected"
    return corrected


# 组装项目可视化数据（前端大屏用）
def _build_project_visualization_payload(
    *,
    identifier_id: str,
    project_detail: dict[str, Any],
    project_result: Optional[dict[str, Any]],
    db_service: PostgreSQLService,
    include_document_content: bool = False,
    include_raw_results: bool = False,
    include_result_record: bool = False,
) -> dict[str, Any]:
    refreshed_result_record, raw_results, display_results = _build_project_display_results(
        identifier_id=identifier_id,
        result_record=project_result,
        db_service=db_service,
    )
    document_identifiers = _collect_project_document_identifiers(project_detail)
    document_map: dict[str, dict[str, Any]] = {}
    for document_identifier in document_identifiers:
        document = db_service.get_document_by_identifier(document_identifier)
        if not document:
            continue
        document_payload = dict(document)
        if not include_document_content:
            document_payload.pop("content", None)
        document_payload["links"] = _document_api_links(document_identifier)
        document_map[document_identifier] = document_payload

    relation_items: list[dict[str, Any]] = []
    for relation in project_detail.get("relations") or []:
        relation_payload = dict(relation)
        relation_payload["documents"] = {}
        for role_name, field_name in (
            ("tender", "tender_identifier_id"),
            ("business_bid", "business_bid_identifier_id"),
            ("technical_bid", "technical_bid_identifier_id"),
        ):
            document_identifier = str(relation.get(field_name) or "").strip()
            if not document_identifier:
                continue
            document_payload = document_map.get(document_identifier)
            if document_payload:
                relation_payload["documents"][role_name] = document_payload
        relation_items.append(relation_payload)

    compact_display_results = _compact_display_results_for_response(display_results)
    payload = {
        "project": project_detail.get("project") or {"identifier_id": identifier_id},
        "project_links": _project_api_links(identifier_id),
        "relations": relation_items,
        "documents": list(document_map.values()),
        "result_record_meta": _build_result_record_meta(refreshed_result_record or project_result),
        "results": compact_display_results,
        "available_result_keys": sorted(compact_display_results.keys()),
    }
    if include_result_record:
        payload["result_record"] = refreshed_result_record or project_result
    if include_raw_results:
        payload["raw_results"] = raw_results
        payload["raw_available_result_keys"] = sorted(raw_results.keys())
    return payload


# 统一持久化上传的 JSON 文件，并返回项目信息与文档记录
async def _persist_uploaded_analysis_documents(
    *,
    tender_json_file: UploadFile,
    business_bid_json_files: Optional[list[UploadFile]],
    technical_bid_json_files: Optional[list[UploadFile]],
    project_name: Optional[str],
    db_service: PostgreSQLService,
) -> tuple[dict, list[dict]]:
    tender_document = await read_uploaded_json_file(
        tender_json_file,
        field_name="tender_json_file",
    )
    business_documents = await load_uploaded_bid_json_documents(
        business_bid_json_files,
        field_name="business_bid_json_files",
        role=DOCUMENT_TYPE_BUSINESS_BID,
    )
    technical_documents = await load_uploaded_bid_json_documents(
        technical_bid_json_files,
        field_name="technical_bid_json_files",
        role=DOCUMENT_TYPE_TECHNICAL_BID,
    )
    persisted_documents = await persist_uploaded_json_project_documents(
        db_service=db_service,
        tender_document=tender_document,
        business_bid_documents=business_documents,
        technical_bid_documents=technical_documents,
        project_name=project_name,
    )
    document_records = build_uploaded_project_document_records(persisted_documents)
    return persisted_documents, document_records


# 将分析结果写入项目结果表
def _persist_uploaded_result(
    *,
    db_service: PostgreSQLService,
    project_identifier: str,
    result_key: str,
    result_value: dict,
) -> dict:
    result = db_service.upsert_project_result_item(
        project_identifier_id=project_identifier,
        result_key=result_key,
        result_value=result_value,
    )
    _invalidate_project_cache_by_identifier(project_identifier)
    return result


# 持久化查重合并结果（多个 key）
def _persist_duplicate_merge_results(
    *,
    db_service: PostgreSQLService,
    project_identifier: str,
    source_result_key: str,
    raw_result: dict[str, Any],
) -> dict[str, dict[str, Any]]:
    merged_results = build_duplicate_merge_results(
        raw_result=raw_result,
        source_result_key=source_result_key,
    )
    for merged_key, merged_value in merged_results.items():
        db_service.upsert_project_result_item(
            project_identifier_id=project_identifier,
            result_key=merged_key,
            result_value=merged_value,
        )
    _invalidate_project_cache_by_identifier(project_identifier)
    return merged_results


# 查找合并结果对应的原始结果键和值
def _resolve_merged_result_source(
    *,
    merged_result_key: str,
    results: dict[str, Any],
) -> tuple[Optional[str], Optional[dict[str, Any]]]:
    doc_type = DOC_TYPE_BY_MERGED_RESULT_KEY.get(merged_result_key)
    if not doc_type:
        return None, None
    preferred_raw_key = RAW_RESULT_KEY_BY_DOC_TYPE.get(doc_type)
    if preferred_raw_key and isinstance(results.get(preferred_raw_key), dict):
        return preferred_raw_key, results.get(preferred_raw_key)
    combined_result = results.get("duplicate_check")
    if isinstance(combined_result, dict):
        return "duplicate_check", combined_result
    return None, None


def _merged_result_needs_rebuild(existing_result: Any) -> bool:
    """派生聚合结果按策略版本重建，避免沿用旧的碎片化 cluster。"""
    if not isinstance(existing_result, dict) or not existing_result:
        return True
    config = existing_result.get("config") or {}
    return str(config.get("merge_strategy") or "").strip() != MERGE_STRATEGY


# 按需加载或重建项目的合并查重结果
def _load_or_build_project_merged_results(
    *,
    identifier_id: str,
    result_record: Optional[dict[str, Any]],
    db_service: PostgreSQLService,
    requested_keys: Optional[list[str]] = None,
) -> tuple[dict[str, Any], dict[str, Any]]:
    results = dict((result_record or {}).get("result") or {})
    target_keys = requested_keys or sorted(
        {
            merged_key
            for merged_key in MERGED_RESULT_KEY_BY_DOC_TYPE.values()
            if (
                merged_key in results
                or _resolve_merged_result_source(merged_result_key=merged_key, results=results)[0]
            )
        }
    )

    merged_payloads: dict[str, Any] = {}
    changed = False
    for merged_key in target_keys:
        existing = results.get(merged_key)
        if not _merged_result_needs_rebuild(existing):
            merged_payloads[merged_key] = existing
            continue

        source_result_key, source_result = _resolve_merged_result_source(
            merged_result_key=merged_key,
            results=results,
        )
        if not source_result_key or not isinstance(source_result, dict):
            continue

        built_results = build_duplicate_merge_results(
            raw_result=source_result,
            source_result_key=source_result_key,
        )
        built_payload = built_results.get(merged_key)
        if not isinstance(built_payload, dict):
            continue

        results[merged_key] = built_payload
        merged_payloads[merged_key] = built_payload
        db_service.upsert_project_result_item(
            project_identifier_id=identifier_id,
            result_key=merged_key,
            result_value=built_payload,
        )
        changed = True

    refreshed_result_record = result_record
    if changed:
        refreshed_result_record = db_service.get_project_result(identifier_id)
        results = dict((refreshed_result_record or {}).get("result") or {})
        merged_payloads = {
            key: value
            for key, value in ((key, results.get(key)) for key in target_keys)
            if isinstance(value, dict)
        }

    return refreshed_result_record or result_record or {}, merged_payloads


# 构建项目展示用结果（含合并查重的替换）
def _build_project_display_results(
    *,
    identifier_id: str,
    result_record: Optional[dict[str, Any]],
    db_service: PostgreSQLService,
) -> tuple[dict[str, Any], dict[str, Any], dict[str, Any]]:
    refreshed_record, merged_results = _load_or_build_project_merged_results(
        identifier_id=identifier_id,
        result_record=result_record,
        db_service=db_service,
    )
    record = refreshed_record or result_record or {}
    result_payload = dict(record.get("result") or {})
    manual_review_results = manual_review_results_from_record(record)
    raw_results = raw_result_view(result_payload)
    display_results = display_result_view(
        result_payload,
        manual_review_results=manual_review_results,
    )
    latest_results = dict(manual_review_results.get("latest") or {})

    merged_key_aliases = {
        "business_bid_duplicate_check": MERGED_RESULT_KEY_BY_DOC_TYPE.get(DOCUMENT_TYPE_BUSINESS_BID),
        "technical_bid_duplicate_check": MERGED_RESULT_KEY_BY_DOC_TYPE.get(DOCUMENT_TYPE_TECHNICAL_BID),
    }
    for raw_key, merged_key in merged_key_aliases.items():
        if not merged_key:
            continue
        if raw_key in latest_results:
            continue
        merged_payload = merged_results.get(merged_key)
        if isinstance(merged_payload, dict) and merged_payload:
            display_results[raw_key] = merged_payload
            display_results[merged_key] = merged_payload

    if "duplicate_check" in display_results:
        display_results["duplicate_check"] = {
            "view_mode": "merged",
            "business_bid_duplicate_check": display_results.get("business_bid_duplicate_check"),
            "technical_bid_duplicate_check": display_results.get("technical_bid_duplicate_check"),
        }

    return refreshed_record or result_record or {}, raw_results, display_results


def _build_workflow_document_summary(payload_data: dict[str, Any], workflow_scope: dict[str, Any]) -> dict[str, Any]:
    all_documents = list(payload_data.get("documents") or [])
    active_payload = filter_project_payload({**payload_data, "workflow_scope": workflow_scope})
    active_documents = list(active_payload.get("documents") or [])
    excluded_count = max(0, len(all_documents) - len(active_documents))

    def _count_role(records: list[dict[str, Any]], role: str) -> int:
        ids = {
            str(record.get("identifier_id") or "").strip()
            for record in records
            if str(record.get("relation_role") or "") == role and str(record.get("identifier_id") or "").strip()
        }
        return len(ids)

    return {
        "total_document_rows": len(all_documents),
        "active_document_rows": len(active_documents),
        "excluded_document_rows": excluded_count,
        "business_bid_count": _count_role(all_documents, DOCUMENT_TYPE_BUSINESS_BID),
        "technical_bid_count": _count_role(all_documents, DOCUMENT_TYPE_TECHNICAL_BID),
        "active_business_bid_count": _count_role(active_documents, DOCUMENT_TYPE_BUSINESS_BID),
        "active_technical_bid_count": _count_role(active_documents, DOCUMENT_TYPE_TECHNICAL_BID),
        "excluded_bidder_count": len(workflow_scope.get(EXCLUDED_BIDDERS_KEY) or []),
    }


def _build_project_workflow_state(
    *,
    identifier_id: str,
    db_service: PostgreSQLService,
) -> dict[str, Any]:
    project = db_service.refresh_project_parsing_status(identifier_id)
    if not project:
        raise HTTPException(status_code=404, detail="project not found")

    payload_data = db_service.get_project_documents_for_duplicate_check(str(project["identifier_id"]))
    if not payload_data:
        raise HTTPException(status_code=404, detail="project not found")

    result_record = db_service.get_project_result(str(project["identifier_id"]))
    manual_review_results = manual_review_results_from_record(result_record)
    workflow_scope = workflow_scope_from_result_record(result_record)
    result_payload = display_result_view(
        (result_record or {}).get("result") or {},
        manual_review_results=manual_review_results,
    )
    result_keys = sorted(
        key
        for key in result_payload.keys()
        if _is_result_key_visible(str(key))
    )
    parsing_status = int(project.get("parsing_status") or 0)
    if parsing_status >= PostgreSQLService.PARSING_STATUS_TECHNICAL_OCR_COMPLETED:
        stage = "technical_ocr_completed"
    elif parsing_status >= PostgreSQLService.PARSING_STATUS_BUSINESS_OCR_COMPLETED:
        stage = "business_review_ready"
    elif parsing_status >= PostgreSQLService.PARSING_STATUS_TENDER_OCR_COMPLETED:
        stage = "business_ocr_pending"
    else:
        stage = "tender_ocr_pending"

    return {
        "project": project,
        "stage": stage,
        "workflow_scope": workflow_scope,
        "excluded_bidders": workflow_scope.get(EXCLUDED_BIDDERS_KEY) or [],
        MANUAL_REVIEW_RESULTS_KEY: manual_review_results,
        "documents": _build_workflow_document_summary(payload_data, workflow_scope),
        "results_status": {
            "available_result_keys": result_keys,
            "business_review_completed": "business_bid_format_review" in result_payload,
            "deviation_check_completed": "deviation_check" in result_payload,
            "business_duplicate_completed": "business_bid_duplicate_check" in result_payload,
            "technical_duplicate_completed": "technical_bid_duplicate_check" in result_payload,
            "personnel_reuse_completed": "personnel_reuse_check" in result_payload,
        },
    }


@router.get("/projects/{identifier_id}/workflow-state", summary="查询项目分步工作台状态")
async def get_project_workflow_state(
    identifier_id: str,
    db_service: PostgreSQLService = Depends(get_db_service),
):
    """返回项目 OCR 阶段、软剔除范围和结果完成状态。"""
    try:
        return await run_in_threadpool(
            _build_project_workflow_state,
            identifier_id=identifier_id,
            db_service=db_service,
        )
    except HTTPException:
        raise
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"database error: {exc}") from exc


@router.put("/projects/{identifier_id}/workflow-scope", summary="保存项目工作流范围")
async def save_project_workflow_scope(
    identifier_id: str,
    payload: ProjectWorkflowScopeRequest,
    db_service: PostgreSQLService = Depends(get_db_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """保存或恢复软剔除投标公司范围。"""
    try:
        workflow_scope = normalize_workflow_scope(payload.model_dump())
        result_record = await run_in_threadpool(
            db_service.update_project_manual_review_workflow_scope,
            project_identifier_id=identifier_id,
            workflow_scope=workflow_scope,
        )
        _invalidate_project_cache_or_error(cache_service, identifier_id)
        manual_review_results = manual_review_results_from_record(result_record)
        saved_scope = workflow_scope_from_result_record(result_record)
        return {
            "project_identifier_id": identifier_id,
            "result_key": WORKFLOW_SCOPE_KEY,
            "workflow_scope": saved_scope,
            "excluded_bidders": saved_scope.get(EXCLUDED_BIDDERS_KEY) or [],
            MANUAL_REVIEW_RESULTS_KEY: manual_review_results,
            "result_record": result_record,
        }
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except HTTPException:
        raise
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"database error: {exc}") from exc


@router.put("/projects/{identifier_id}/manual-review-results/{result_key}/inputs", summary="保存最新人工审查结果输入")
async def save_project_manual_review_result_inputs(
    identifier_id: str,
    result_key: str,
    payload: ProjectManualReviewResultInputsRequest,
    db_service: PostgreSQLService = Depends(get_db_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """保存任意结果键的人工判断补丁，结果写入 manual_review_results.latest。"""
    normalized_result_key = str(result_key or "").strip()
    if not normalized_result_key:
        raise HTTPException(status_code=400, detail="result_key is required")
    try:
        current_record = await run_in_threadpool(db_service.get_project_result, identifier_id)
        manual_review_results = manual_review_results_from_record(current_record)
        result_payload = dict((current_record or {}).get("result") or {})
        latest_payload = dict((manual_review_results.get("latest") or {}).get(normalized_result_key) or {})
        base_payload = latest_payload or dict(result_payload.get(normalized_result_key) or {})
        if not base_payload:
            raise HTTPException(status_code=404, detail=f"result key not found: {normalized_result_key}")
        next_result_value = _apply_manual_result_inputs(base_payload, payload.inputs)
        result_record = await run_in_threadpool(
            db_service.update_project_manual_review_result,
            project_identifier_id=identifier_id,
            result_key=normalized_result_key,
            result_value=next_result_value,
        )
        _invalidate_project_cache_or_error(cache_service, identifier_id)
        manual_review_results = manual_review_results_from_record(result_record)
        return {
            "project_identifier_id": identifier_id,
            "result_key": normalized_result_key,
            MANUAL_REVIEW_RESULTS_KEY: manual_review_results,
            "latest": (manual_review_results.get("latest") or {}).get(normalized_result_key) or {},
            "result_record": result_record,
        }
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except HTTPException:
        raise
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"database error: {exc}") from exc


@router.post("/projects/{identifier_id}/manual-review-rerun", summary="人工修订后项目级重审")
async def rerun_project_manual_review_results(
    identifier_id: str,
    payload: ProjectManualReviewRerunRequest,
    db_service: PostgreSQLService = Depends(get_db_service),
    duplicate_check_service: DuplicateCheckService = Depends(get_duplicate_check_service),
    bid_document_review_service: BidDocumentReviewService = Depends(get_bid_document_review_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """Run selected services against effective document content and save outputs as latest manual results."""
    selected_services = [
        service
        for service in _normalize_selected_services(payload.services)
        if service
    ]
    if not selected_services:
        raise HTTPException(status_code=400, detail="services cannot be empty")

    latest_results: dict[str, Any] = {}
    items: list[dict[str, Any]] = []
    try:
        project = await run_in_threadpool(_refresh_project_or_404, db_service, identifier_id)
        _ensure_project_ocr_idle(project, analysis_name="人工重审")

        for service_name in selected_services:
            result_key = _PROJECT_SERVICE_RESULT_KEYS.get(service_name, service_name)
            try:
                if service_name == "business_bid_format_review":
                    _ensure_project_analysis_status(
                        project,
                        required_status=PostgreSQLService.PARSING_STATUS_BUSINESS_OCR_COMPLETED,
                        analysis_name="商务标形式重审",
                    )
                    review_service = UnifiedBusinessReviewService(db_service=db_service)
                    rerun_result = await run_in_threadpool(
                        review_service.persist_project_business_review,
                        project_identifier=identifier_id,
                        result_key=BUSINESS_FORMAT_RESULT_KEY,
                    )
                    review = rerun_result.get("review") or {}
                    manual_payload = await run_in_threadpool(
                        _business_manual_payload_for_project,
                        identifier_id=identifier_id,
                        db_service=db_service,
                    )
                    result_value = _apply_manual_business_review_inputs(review, manual_payload)
                    await run_in_threadpool(
                        db_service.update_project_manual_review_result,
                        project_identifier_id=identifier_id,
                        result_key=BUSINESS_FORMAT_RESULT_KEY,
                        result_value=result_value,
                    )
                elif service_name == "deviation_check":
                    _ensure_project_analysis_status(
                        project,
                        required_status=PostgreSQLService.PARSING_STATUS_TECHNICAL_OCR_COMPLETED,
                        analysis_name="偏离表重审",
                    )
                    result_payload = await run_in_threadpool(
                        _run_project_deviation_check,
                        identifier_id=identifier_id,
                        db_service=db_service,
                        persist_to_latest=True,
                    )
                    result_value = result_payload.get("review") or {}
                elif service_name == "business_bid_duplicate_check":
                    _ensure_project_analysis_status(
                        project,
                        required_status=PostgreSQLService.PARSING_STATUS_BUSINESS_OCR_COMPLETED,
                        analysis_name="商务标查重重审",
                    )
                    result_value = await run_in_threadpool(
                        _run_project_duplicate_check,
                        identifier_id=identifier_id,
                        document_types=["business_bid"],
                        max_evidence_sections=5,
                        max_pairs_per_type=0,
                        result_key="business_bid_duplicate_check",
                        db_service=db_service,
                        duplicate_check_service=duplicate_check_service,
                        persist_to_latest=True,
                    )
                elif service_name == "technical_bid_duplicate_check":
                    _ensure_project_analysis_status(
                        project,
                        required_status=PostgreSQLService.PARSING_STATUS_TECHNICAL_OCR_COMPLETED,
                        analysis_name="技术标查重重审",
                    )
                    result_value = await run_in_threadpool(
                        _run_project_duplicate_check,
                        identifier_id=identifier_id,
                        document_types=["technical_bid"],
                        max_evidence_sections=5,
                        max_pairs_per_type=0,
                        result_key="technical_bid_duplicate_check",
                        db_service=db_service,
                        duplicate_check_service=duplicate_check_service,
                        persist_to_latest=True,
                    )
                elif service_name == "personnel_reuse_check":
                    _ensure_project_analysis_status(
                        project,
                        required_status=PostgreSQLService.PARSING_STATUS_TECHNICAL_OCR_COMPLETED,
                        analysis_name="人员复用重审",
                    )
                    result_value = await run_in_threadpool(
                        _run_project_personnel_reuse_check,
                        identifier_id=identifier_id,
                        db_service=db_service,
                        bid_document_review_service=bid_document_review_service,
                        persist_to_latest=True,
                    )
                else:
                    raise HTTPException(status_code=400, detail=f"不支持的人工重审服务：{service_name}")

                latest_results[result_key] = result_value
                items.append({"service": service_name, "result_key": result_key, "status": "completed"})
            except HTTPException as exc:
                items.append(
                    {
                        "service": service_name,
                        "result_key": result_key,
                        "status": "failed",
                        "status_code": exc.status_code,
                        "error": _http_exception_detail_to_text(exc.detail),
                    }
                )

        result_record = await run_in_threadpool(db_service.get_project_result, identifier_id)
        manual_review_results = manual_review_results_from_record(result_record)
        _invalidate_project_cache_or_error(cache_service, identifier_id)
        return {
            "project_identifier_id": identifier_id,
            "services": selected_services,
            "items": items,
            "latest": manual_review_results.get("latest") or latest_results,
            MANUAL_REVIEW_RESULTS_KEY: manual_review_results,
            "result_record": result_record,
        }
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"database error: {exc}") from exc


_REPORT_RESULT_LABELS: tuple[tuple[str, str], ...] = (
    ("business_bid_duplicate_check", "商务标查重"),
    ("technical_bid_duplicate_check", "技术标查重"),
    ("business_bid_format_review", "商务标形式审查"),
    ("deviation_check", "偏离表检查"),
    ("personnel_reuse_check", "人员复用检查"),
    ("business_bid_duplicate_clusters", "商务标查重聚类"),
    ("technical_bid_duplicate_clusters", "技术标查重聚类"),
)
_REPORT_RESULT_LABEL_BY_KEY = dict(_REPORT_RESULT_LABELS)
_REPORT_ISSUE_TABLE_HEADERS = ["问题项目", "对应文件", "页码", "问题描述"]
_REPORT_WORD_LATIN_FONT = "Times New Roman"
_REPORT_WORD_EAST_ASIA_FONT = "宋体"


def _report_compact_text(value: Any, *, max_length: int = 260) -> str:
    if value is None:
        text = ""
    elif isinstance(value, str):
        text = value
    elif isinstance(value, (int, float, bool)):
        text = str(value)
    else:
        try:
            text = json.dumps(value, ensure_ascii=False, separators=(",", ":"))
        except TypeError:
            text = str(value)
    return re.sub(r"\s+", " ", text).strip()


def _report_page_text(item: dict[str, Any], *, side: str | None = None) -> str:
    page_keys = []
    if side:
        page_keys.extend([f"{side}_source_page", f"{side}_page", f"{side}_pages"])
    page_keys.extend(["source_page", "page", "pages", "start_page"])
    for key in page_keys:
        value = item.get(key)
        if value not in (None, "", []):
            return _report_compact_text(value, max_length=120)
    return ""


def _report_file_name(item: dict[str, Any], *, side: str | None = None) -> str:
    name_keys = []
    if side:
        name_keys.extend([f"{side}_file_name", f"{side}_document_name"])
    name_keys.extend(["file_name", "document_name", "name"])
    for key in name_keys:
        value = str(item.get(key) or "").strip()
        if value:
            return value
    return ""


def _report_issue_title(issue: dict[str, Any]) -> str:
    for key in ("title", "check_name", "display_text", "highlight_text", "matched_text", "name", "issue_type", "risk_level"):
        value = str(issue.get(key) or "").strip()
        if value:
            return value
    return _report_compact_text(issue, max_length=80)


def _report_issue_message(issue: dict[str, Any]) -> str:
    for key in ("reason", "message", "summary", "text", "suggestion", "preview", "description", "detail"):
        value = issue.get(key)
        if value not in (None, "", []):
            return _report_compact_text(value)
    return ""


def _report_issue_description(issue: dict[str, Any]) -> str:
    parts: list[Any] = []
    field_labels = (
        ("display_text", "展示文本"),
        ("raw_matched_text", "原始命中"),
        ("model_matched_text", "模型命中"),
        ("matched_text", "命中文本"),
        ("wrong", "疑似文本"),
        ("suggestion", "建议"),
        ("correct", "建议"),
        ("context", "上下文"),
        ("preview", "片段"),
        ("description", "描述"),
        ("detail", "详情"),
        ("message", "说明"),
        ("text", "文本"),
    )
    for key, label in field_labels:
        value = issue.get(key)
        if value not in (None, "", []):
            parts.append(f"{label}：{_report_compact_text(value, max_length=300)}")
    return _report_join_parts(parts, max_length=500)


def _report_join_parts(parts: list[Any], *, max_length: int = 500) -> str:
    texts: list[str] = []
    for part in parts:
        text = _report_compact_text(part, max_length=max_length).strip()
        if text and text not in texts:
            texts.append(text)
    return "；".join(texts)


def _report_file_and_page(item: dict[str, Any], *, side: str | None = None) -> tuple[str, str]:
    file_name = _report_file_name(item, side=side)
    page = _report_page_text(item, side=side)
    if not side and isinstance(item.get("locations"), list):
        location_files: list[str] = []
        location_pages: list[str] = []
        for location in item.get("locations") or []:
            if not isinstance(location, dict):
                continue
            location_file = str(location.get("file_name") or "").strip()
            location_page = _report_compact_text(location.get("page"), max_length=40)
            if location_file and location_file not in location_files:
                location_files.append(location_file)
            if location_page and location_page not in location_pages:
                location_pages.append(location_page)
        file_name = file_name or "; ".join(location_files)
        page = page or "; ".join(location_pages)
    return file_name, page


def _report_ranges_by_file(value: Any) -> str:
    if not isinstance(value, dict):
        return ""
    parts: list[str] = []
    for file_name, ranges in value.items():
        range_parts: list[str] = []
        if isinstance(ranges, list):
            for item in ranges:
                if isinstance(item, dict):
                    start = item.get("start_page")
                    end = item.get("end_page")
                    if start and end and start != end:
                        range_parts.append(f"{start}-{end}")
                    elif start:
                        range_parts.append(str(start))
                elif item not in (None, ""):
                    range_parts.append(str(item))
        if range_parts:
            parts.append(f"{file_name}: {', '.join(range_parts)}")
    return "; ".join(parts)


def _ordered_report_results(result_payload: dict[str, Any]) -> dict[str, Any]:
    ordered: dict[str, Any] = {}
    for key, _label in _REPORT_RESULT_LABELS:
        if key in result_payload and _is_result_key_visible(key):
            ordered[key] = result_payload[key]
    for key, value in result_payload.items():
        if key not in ordered and _is_result_key_visible(key):
            ordered[key] = value
    return ordered


def _set_word_rfonts(target: Any) -> None:
    rpr = target.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), _REPORT_WORD_LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), _REPORT_WORD_LATIN_FONT)
    rfonts.set(qn("w:cs"), _REPORT_WORD_LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), _REPORT_WORD_EAST_ASIA_FONT)


def _set_word_run_fonts(run: Any) -> None:
    run.font.name = _REPORT_WORD_LATIN_FONT
    _set_word_rfonts(run._element)


def _set_word_style_fonts(style: Any) -> None:
    style.font.name = _REPORT_WORD_LATIN_FONT
    _set_word_rfonts(style._element)


def _apply_report_document_fonts(document: Document) -> None:
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet"):
        try:
            _set_word_style_fonts(document.styles[style_name])
        except KeyError:
            continue

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            _set_word_run_fonts(run)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _set_word_run_fonts(run)


def _add_word_paragraph(document: Document, text: str, *, style: Optional[str] = None) -> None:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    _set_word_run_fonts(run)


def _set_word_cell_text(cell: Any, text: Any, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text or ""))
    run.bold = bold
    _set_word_run_fonts(run)


def _add_word_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    if not headers:
        return
    table = document.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _set_word_cell_text(cell, header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        normalized = list(row[: len(headers)])
        if len(normalized) < len(headers):
            normalized.extend([""] * (len(headers) - len(normalized)))
        for index, value in enumerate(normalized):
            _set_word_cell_text(cells[index], value)
    document.add_paragraph()


def _report_issue_table_row(row: list[str]) -> list[str]:
    normalized = list(row[:5])
    if len(normalized) < 5:
        normalized.extend([""] * (5 - len(normalized)))
    problem, description, _reason, file_name, page = normalized
    return [problem, file_name, page, description]


def _add_report_issue_table(document: Document, rows: list[list[str]]) -> None:
    _add_word_table(
        document,
        _REPORT_ISSUE_TABLE_HEADERS,
        [_report_issue_table_row(row) for row in rows],
    )


def _report_project_name(project: Optional[dict[str, Any]]) -> str:
    if not isinstance(project, dict):
        return ""
    return str(project.get("project_name") or project.get("identifier_id") or "").strip()


def _report_bound_document_rows(project_detail: Optional[dict[str, Any]]) -> list[list[str]]:
    field_groups = (
        ("招标文件", "tender_identifier_id", "tender_file_name"),
        ("商务标文件", "business_bid_identifier_id", "business_bid_file_name"),
        ("技术标文件", "technical_bid_identifier_id", "technical_bid_file_name"),
    )
    rows: list[list[str]] = []
    seen: set[tuple[str, str]] = set()
    for relation in (project_detail or {}).get("relations") or []:
        if not isinstance(relation, dict):
            continue
        for label, identifier_field, file_name_field in field_groups:
            identifier = str(relation.get(identifier_field) or "").strip()
            file_name = str(relation.get(file_name_field) or "").strip()
            if not identifier and not file_name:
                continue
            key = (label, identifier or file_name)
            if key in seen:
                continue
            seen.add(key)
            rows.append([label, file_name or identifier])
    return rows


def _report_issue_row(
    rows: list[list[str]],
    *,
    problem: Any,
    reason: Any,
    description: Any = "",
    file_name: Any = "",
    page: Any = "",
) -> None:
    problem_text = _report_compact_text(problem, max_length=260)
    description_text = _report_compact_text(description, max_length=700)
    reason_text = _report_compact_text(reason, max_length=500)
    file_text = _report_compact_text(file_name, max_length=260)
    page_text = _report_compact_text(page, max_length=120)
    if not any((problem_text, description_text, reason_text, file_text, page_text)):
        return
    rows.append([
        problem_text or "未命名问题",
        description_text or "-",
        reason_text or "-",
        file_text or "-",
        page_text or "-",
    ])


def _collect_personnel_issue_rows(payload: dict[str, Any]) -> list[list[str]]:
    rows: list[list[str]] = []
    groups = payload.get("groups") or {"default": payload}
    for group in groups.values():
        if not isinstance(group, dict):
            continue
        check = (
            group.get("personnel_reuse_check")
            if isinstance(group.get("personnel_reuse_check"), dict)
            else group
        )
        for item in check.get("issues") or []:
            if not isinstance(item, dict):
                continue
            evidence_items = item.get("occurrences") if isinstance(item.get("occurrences"), list) else [item]
            reason = _report_issue_message(item) or _report_join_parts([
                item.get("risk_level"),
                item.get("roles"),
                f"涉及文件数：{item.get('document_count')}" if item.get("document_count") else "",
                f"出现次数：{item.get('occurrence_count')}" if item.get("occurrence_count") else "",
            ])
            description = _report_join_parts([
                f"姓名：{item.get('name')}" if item.get("name") else "",
                f"角色：{item.get('roles')}" if item.get("roles") else "",
                f"涉及文件数：{item.get('document_count')}" if item.get("document_count") else "",
                f"出现次数：{item.get('occurrence_count')}" if item.get("occurrence_count") else "",
            ])
            for evidence in evidence_items or [item]:
                if not isinstance(evidence, dict):
                    continue
                file_name, page = _report_file_and_page(evidence)
                _report_issue_row(
                    rows,
                    problem=item.get("name") or _report_issue_title(item) or "人员复用",
                    description=description or _report_issue_description(evidence),
                    reason=reason or "疑似一人多用",
                    file_name=file_name or _report_file_name(item),
                    page=page or _report_page_text(item),
                )
    return rows


def _collect_business_format_issue_rows(payload: dict[str, Any]) -> list[list[str]]:
    rows: list[list[str]] = []
    frontend_items = payload.get("issues") if isinstance(payload.get("issues"), list) else []
    if frontend_items:
        for item in frontend_items:
            if not isinstance(item, dict):
                continue
            issue = item.get("issue") if isinstance(item.get("issue"), dict) else item
            file_name, page = _report_file_and_page(item)
            issue_file, issue_page = _report_file_and_page(issue)
            _report_issue_row(
                rows,
                problem=_report_join_parts([
                    item.get("bidder_name"),
                    item.get("check_name"),
                    _report_issue_title(issue),
                ]),
                description=_report_issue_description(issue) or _report_issue_message(issue),
                reason=_report_join_parts([
                    item.get("check_code"),
                    issue.get("severity"),
                    issue.get("status"),
                ]),
                file_name=file_name or issue_file,
                page=page or issue_page,
            )
        return rows

    for bidder in payload.get("bidders") or []:
        if not isinstance(bidder, dict):
            continue
        bidder_name = bidder.get("bidder_name") or bidder.get("bidder_key")
        for status_key, issues in (bidder.get("issues") or {}).items():
            if not isinstance(issues, list):
                continue
            for issue in issues:
                if not isinstance(issue, dict):
                    continue
                file_name, page = _report_file_and_page(issue)
                _report_issue_row(
                    rows,
                    problem=_report_join_parts([bidder_name, _report_issue_title(issue)]),
                    description=_report_issue_description(issue),
                    reason=_report_join_parts([status_key, issue.get("severity"), _report_issue_message(issue)]),
                    file_name=file_name,
                    page=page,
                )
    return rows


def _collect_duplicate_issue_rows(payload: dict[str, Any]) -> list[list[str]]:
    rows: list[list[str]] = []
    if isinstance(payload.get("issues"), list) and not payload.get("groups"):
        return _collect_duplicate_cluster_issue_rows(payload)

    for group in (payload.get("groups") or {}).values():
        if not isinstance(group, dict):
            continue
        for item in group.get("issues") or []:
            if not isinstance(item, dict):
                continue
            evidence_parts: list[str] = []
            for block in (item.get("duplicate_blocks") or item.get("evidence_sections") or [])[:3]:
                if isinstance(block, dict):
                    evidence_parts.append(_report_issue_message(block))
            left_file, left_page = _report_file_and_page(item, side="left")
            right_file, right_page = _report_file_and_page(item, side="right")
            _report_issue_row(
                rows,
                problem=_report_issue_title(item) or "重复内容",
                description=_report_join_parts([
                    _report_issue_description(item),
                    *evidence_parts,
                ]),
                reason=_report_join_parts([
                    item.get("risk_level"),
                    item.get("score_display") or item.get("similarity"),
                    _report_issue_message(item),
                ]),
                file_name=_report_join_parts([
                    f"左：{left_file}" if left_file else "",
                    f"右：{right_file}" if right_file else "",
                    _report_file_name(item),
                ]),
                page=_report_join_parts([
                    f"左：{left_page}" if left_page else "",
                    f"右：{right_page}" if right_page else "",
                    _report_page_text(item),
                ], max_length=180),
            )
    return rows


def _collect_duplicate_cluster_issue_rows(payload: dict[str, Any]) -> list[list[str]]:
    rows: list[list[str]] = []
    for cluster in payload.get("issues") or []:
        if not isinstance(cluster, dict):
            continue
        _report_issue_row(
            rows,
            problem=cluster.get("title") or cluster.get("cluster_id") or "重复聚类",
            description=_report_join_parts([
                f"聚类ID：{cluster.get('cluster_id')}" if cluster.get("cluster_id") else "",
                f"文件：{cluster.get('files')}" if cluster.get("files") else "",
                f"页码：{_report_ranges_by_file(cluster.get('doc_ranges_by_file'))}" if cluster.get("doc_ranges_by_file") else "",
            ]),
            reason=_report_join_parts([
                cluster.get("risk_level"),
                cluster.get("score_display") or cluster.get("similarity"),
                f"文件数：{cluster.get('file_count')}" if cluster.get("file_count") else "",
                f"出现次数：{cluster.get('occurrence_count')}" if cluster.get("occurrence_count") else "",
            ]),
            file_name=cluster.get("files"),
            page=_report_ranges_by_file(cluster.get("doc_ranges_by_file")),
        )
    return rows


def _collect_generic_issue_rows(section_key: str, payload: Any) -> list[list[str]]:
    rows: list[list[str]] = []
    label = _REPORT_RESULT_LABEL_BY_KEY.get(section_key, section_key)

    def visit(value: Any, parent_key: str = "") -> None:
        if isinstance(value, list):
            for item in value:
                visit(item, parent_key)
            return
        if not isinstance(value, dict):
            if parent_key in {"issues", "failed", "missing", "unclear"}:
                _report_issue_row(rows, problem=label, reason=value)
            return

        has_issue_key = any(
            key in value
            for key in (
                "title",
                "check_name",
                "matched_text",
                "issue_type",
                "risk_level",
                "reason",
                "message",
                "suggestion",
            )
        )
        has_container_key = any(
            key in value
            for key in ("groups", "documents", "bidders", "checks", "issues", "summary")
        )
        if has_issue_key and not has_container_key:
            file_name, page = _report_file_and_page(value)
            _report_issue_row(
                rows,
                problem=_report_issue_title(value) or label,
                description=_report_issue_description(value),
                reason=_report_issue_message(value) or label,
                file_name=file_name,
                page=page,
            )
            return

        for key, child in value.items():
            if key == "summary":
                continue
            visit(child, key)

    visit(payload)
    return rows


def _collect_report_section_issue_rows(section_key: str, payload: Any) -> list[list[str]]:
    collectors = {
        "personnel_reuse_check": _collect_personnel_issue_rows,
        "business_bid_format_review": _collect_business_format_issue_rows,
        "business_bid_duplicate_check": _collect_duplicate_issue_rows,
        "technical_bid_duplicate_check": _collect_duplicate_issue_rows,
        "business_bid_duplicate_clusters": _collect_duplicate_cluster_issue_rows,
        "technical_bid_duplicate_clusters": _collect_duplicate_cluster_issue_rows,
    }
    section_rows: list[list[str]] = []
    if isinstance(payload, dict) and section_key in collectors:
        section_rows = collectors[section_key](payload)
    if not section_rows:
        section_rows = _collect_generic_issue_rows(section_key, payload)
    return section_rows


def _dedupe_report_issue_rows(rows: list[list[str]]) -> list[list[str]]:
    deduped: list[list[str]] = []
    seen: set[tuple[str, str, str, str, str]] = set()
    for row in rows:
        normalized = tuple(row[:5])
        if normalized in seen:
            continue
        seen.add(normalized)
        deduped.append(row)
    return deduped


def _collect_frontend_result_issue_sections(items: list[Any]) -> list[tuple[str, str, list[list[str]]]]:
    grouped: dict[str, list[dict[str, Any]]] = {}
    order: list[str] = []
    for item in items:
        if not isinstance(item, dict):
            continue
        key = str(item.get("result_key") or item.get("source_result_key") or item.get("type") or "result").strip()
        if not key:
            key = "result"
        if not _is_result_key_visible(key):
            continue
        if key not in grouped:
            grouped[key] = []
            order.append(key)
        grouped[key].append(item)

    sections: list[tuple[str, str, list[list[str]]]] = []
    for key in order:
        section_rows = _dedupe_report_issue_rows(
            _collect_report_section_issue_rows(key, {"issues": grouped[key]})
        )
        if not section_rows:
            continue
        sections.append((key, _REPORT_RESULT_LABEL_BY_KEY.get(key, "自定义审查项"), section_rows))
    return sections


def _collect_report_issue_sections(result_payload: dict[str, Any]) -> list[tuple[str, str, list[list[str]]]]:
    sections: list[tuple[str, str, list[list[str]]]] = []
    frontend_items = result_payload.get("result") if isinstance(result_payload, dict) else None
    if isinstance(frontend_items, list):
        return _collect_frontend_result_issue_sections(frontend_items)

    for key, payload in _ordered_report_results(result_payload).items():
        section_rows = _dedupe_report_issue_rows(_collect_report_section_issue_rows(key, payload))
        if not section_rows:
            continue
        sections.append((key, _REPORT_RESULT_LABEL_BY_KEY.get(key, "自定义审查项"), section_rows))
    return sections


def _flatten_report_issue_sections(
    sections: list[tuple[str, str, list[list[str]]]],
) -> list[list[str]]:
    rows: list[list[str]] = []
    for _key, _label, section_rows in sections:
        rows.extend(section_rows)
    return _dedupe_report_issue_rows(rows)


def _collect_report_issue_rows(result_payload: dict[str, Any]) -> list[list[str]]:
    return _flatten_report_issue_sections(_collect_report_issue_sections(result_payload))


def _render_result_word_report(
    result_payload: dict[str, Any],
    *,
    project: Optional[dict[str, Any]] = None,
    project_detail: Optional[dict[str, Any]] = None,
    exported_at: Optional[datetime] = None,
) -> bytes:
    document = Document()
    resolved_project = project or (project_detail or {}).get("project") or {}
    issue_sections = _collect_report_issue_sections(result_payload)
    document.add_heading("项目审查报告", level=0)
    _add_word_table(
        document,
        ["项目", "内容"],
        [
            ["导出时间", (exported_at or datetime.now()).isoformat(timespec="seconds")],
            ["导出项目", _report_project_name(resolved_project) or "-"],
            ["审查项数量", str(len(issue_sections))],
        ],
    )

    document.add_heading("总览", level=1)
    if issue_sections:
        _add_word_table(
            document,
            ["审查项", "问题数"],
            [[label, str(len(section_rows))] for _key, label, section_rows in issue_sections],
        )
    else:
        _add_word_paragraph(document, "无")

    document.add_heading("对应文件", level=1)
    document_rows = _report_bound_document_rows(project_detail)
    if document_rows:
        _add_word_table(document, ["文件类型", "文件名称"], document_rows)
    else:
        _add_word_paragraph(document, "无")

    document.add_heading("问题清单", level=1)
    if issue_sections:
        for _key, label, section_rows in issue_sections:
            document.add_heading(label, level=2)
            _add_report_issue_table(document, section_rows)
    else:
        _add_word_paragraph(document, "无")

    _apply_report_document_fonts(document)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _safe_report_file_name(raw_name: Optional[str]) -> str:
    name = str(raw_name or "project_result_report").strip()
    name = os.path.splitext(os.path.basename(name))[0]
    name = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff._-]+", "_", name).strip("._-")
    return (name or "project_result_report")[:80]


async def _upload_word_report(
    *,
    report_bytes: bytes,
    report_name: Optional[str],
    oss_service: MinioService,
    object_prefix: str = "report",
) -> dict[str, Any]:
    safe_name = _safe_report_file_name(report_name)
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S%f")
    digest = hashlib.sha1(report_bytes).hexdigest()[:8]
    prefix = str(object_prefix or "report").strip().strip("/") or "report"
    object_name = f"{prefix}/{safe_name}_{timestamp}_{digest}.docx"
    return await run_in_threadpool(
        oss_service.upload_bytes,
        report_bytes,
        filename=f"{safe_name}.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        object_name=object_name,
    )


async def _upload_result_word_report(
    *,
    project: dict[str, Any],
    project_detail: Optional[dict[str, Any]] = None,
    result_payload: dict[str, Any],
    oss_service: MinioService,
) -> dict[str, Any]:
    report_name = f"{project.get('project_name') or project['identifier_id']}_最终导出报告"
    report_bytes = _render_result_word_report(
        result_payload,
        project=project,
        project_detail=project_detail,
    )
    upload = await _upload_word_report(
        report_bytes=report_bytes,
        report_name=report_name,
        oss_service=oss_service,
        object_prefix="report",
    )
    return {
        "report_name": f"{_safe_report_file_name(report_name)}.docx",
        "upload": upload,
    }


def _attach_report_to_response(
    *,
    response: dict[str, Any],
    project: dict[str, Any],
    report: dict[str, Any],
    db_service: PostgreSQLService,
) -> dict[str, Any]:
    upload = dict(report.get("upload") or {})
    report_url = str(upload.get("file_url") or "").strip()
    updated_project = None
    if report_url:
        updated_project = db_service.update_project_report_url(
            str(project["identifier_id"]),
            report_url,
        )

    response["report_upload"] = upload
    response["report_name"] = report.get("report_name")
    response["report_url"] = report_url
    if updated_project:
        response["project"] = updated_project
    return response


async def _export_result_word_report(
    *,
    project_identifier_id: str,
    result_payload: dict[str, Any],
    db_service: PostgreSQLService,
    oss_service: MinioService,
) -> dict[str, Any]:
    project_detail = db_service.get_project_detail(project_identifier_id)
    if not project_detail:
        raise ValueError(f"项目不存在：{project_identifier_id}")
    project = project_detail.get("project") or {}

    report = await _upload_result_word_report(
        project=project,
        project_detail=project_detail,
        result_payload=result_payload,
        oss_service=oss_service,
    )
    response = {"project_identifier_id": str(project["identifier_id"])}
    return _attach_report_to_response(
        response=response,
        project=project,
        report=report,
        db_service=db_service,
    )


# ======================== 路由定义 ========================

# 项目 CRUD
@router.post("/projects", summary="创建项目")
async def create_project(
    payload: ProjectCreateRequest,
    db_service: PostgreSQLService = Depends(get_db_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """创建新项目，项目 UUID 由系统自动生成，若项目名称已存在则返回 409。"""
    try:
        created = db_service.create_project(
            project_name=payload.project_name,
        )
        _invalidate_project_cache_or_error(cache_service)
        return created
    except PsycopgError as exc:
        if getattr(exc, "pgcode", None) == "23505":
            raise HTTPException(status_code=409, detail="项目名称已存在") from exc
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.get("/projects", summary="查询项目列表")
async def list_projects(
    response: Response,
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=20, ge=1, le=200),
    limit: Optional[int] = Query(default=None, ge=1, le=200),
    offset: Optional[int] = Query(default=None, ge=0),
    keyword: Optional[str] = Query(default=None, description="按项目名称或 UUID 模糊搜索"),
    db_service: PostgreSQLService = Depends(get_db_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """分页查询项目列表，支持关键字搜索。"""
    try:
        resolved_limit, resolved_offset = _resolve_pagination(
            page=page,
            page_size=page_size,
            limit=limit,
            offset=offset,
        )
        cache_key = cache_service.project_list_key(
            limit=resolved_limit,
            offset=resolved_offset,
            keyword=keyword,
        )
        return _cache_get_or_set_payload(
            cache_service=cache_service,
            cache_key=cache_key,
            ttl_seconds=settings.XTJS_CACHE_PROJECT_LIST_TTL_SECONDS,
            response=response,
            factory=lambda: db_service.list_projects(
                limit=resolved_limit,
                offset=resolved_offset,
                keyword=keyword,
            ),
        )
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.post("/projects/batch-delete", summary="批量删除项目")
async def batch_delete_projects(
    payload: IdentifierBatchDeleteRequest,
    db_service: PostgreSQLService = Depends(get_db_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """软删除指定标识集合中的项目。"""
    try:
        deleted_count = db_service.soft_delete_projects(payload.identifier_ids)
        for identifier_id in payload.identifier_ids:
            if str(identifier_id or "").strip():
                cancel_project_runtime(identifier_id)
        _invalidate_project_cache_or_error(cache_service)
        return {
            "requested_count": len(payload.identifier_ids),
            "deleted_count": deleted_count,
        }
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.get("/projects/{identifier_id}", summary="查询项目详情")
async def get_project_detail(
    identifier_id: str,
    response: Response,
    db_service: PostgreSQLService = Depends(get_db_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """按项目标识返回项目详情（含绑定关系）。"""
    try:
        cache_key = cache_service.project_detail_key(identifier_id)

        def _load_detail():
            detail = db_service.get_project_detail(identifier_id)
            if not detail:
                raise HTTPException(status_code=404, detail="项目不存在")
            return detail

        return _cache_get_or_set_payload(
            cache_service=cache_service,
            cache_key=cache_key,
            ttl_seconds=settings.XTJS_CACHE_PROJECT_DETAIL_TTL_SECONDS,
            response=response,
            factory=_load_detail,
        )
    except HTTPException:
        raise
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.put("/projects/{identifier_id}", summary="更新项目名称")
async def update_project(
    identifier_id: str,
    payload: ProjectUpdateRequest,
    db_service: PostgreSQLService = Depends(get_db_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """修改项目名称，项目 UUID 不允许修改。"""
    try:
        updated = db_service.update_project(
            identifier_id=identifier_id,
            project_name=payload.project_name,
        )
        if not updated:
            raise HTTPException(status_code=404, detail="项目不存在")
        _invalidate_project_cache_or_error(cache_service, str(updated.get("identifier_id") or identifier_id))
        return updated
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        if getattr(exc, "pgcode", None) == "23505":
            raise HTTPException(status_code=409, detail="项目名称已存在") from exc
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.delete("/projects/{identifier_id}", summary="删除项目")
async def delete_project(
    identifier_id: str,
    db_service: PostgreSQLService = Depends(get_db_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """软删除项目。"""
    try:
        existing_project = db_service.get_project_by_identifier(identifier_id)
        if not existing_project:
            raise HTTPException(status_code=404, detail="项目不存在")
        resolved_identifier = str(existing_project["identifier_id"])
        deleted = db_service.soft_delete_project(identifier_id)
        if not deleted:
            raise HTTPException(status_code=404, detail="项目不存在")
        cancel_project_runtime(resolved_identifier)
        _invalidate_project_cache_or_error(cache_service, resolved_identifier)
        _invalidate_project_cache_or_error(cache_service)
        return {"status": "deleted"}
    except HTTPException:
        raise
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


# 项目分析结果
@router.get("/projects/{identifier_id}/results", summary="查询项目分析结果")
async def get_project_results(
    identifier_id: str,
    response: Response,
    view: Literal["display", "raw"] = Query(
        default="display",
        description="display=默认返回 merge 后的展示结果，raw=返回原始结果",
    ),
    include_raw_results: bool = Query(default=False),
    include_result_record: bool = Query(default=False),
    force_refresh: bool = Query(default=False),
    db_service: PostgreSQLService = Depends(get_db_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """获取项目的分析结果，支持展示视图和原始视图，可附加返回原始数据。"""
    try:
        cache_key = cache_service.project_results_key(
            identifier_id,
            view=view,
            include_raw_results=include_raw_results,
            include_result_record=include_result_record,
        )

        def _load_results():
            project = db_service.get_project_by_identifier(identifier_id)
            if not project:
                raise HTTPException(status_code=404, detail="project not found")

            result_record = db_service.get_project_result(identifier_id)
            refreshed_record = result_record or {}
            raw_results = raw_result_view((result_record or {}).get("result") or {})
            display_results = raw_results
            if view == "display" or include_raw_results:
                refreshed_record, raw_results, display_results = _build_project_display_results(
                    identifier_id=identifier_id,
                    result_record=result_record,
                    db_service=db_service,
                )

            selected_results = (
                _compact_display_results_for_response(display_results)
                if view == "display"
                else raw_results
            )
            selected_result_keys = sorted(selected_results.keys())
            payload = {
                "project": project,
                "view": view,
                "result_record_meta": _build_result_record_meta(refreshed_record or result_record),
                "results": selected_results,
                "available_result_keys": selected_result_keys,
            }
            if include_result_record:
                payload["result_record"] = refreshed_record or result_record
            if include_raw_results:
                payload["raw_results"] = raw_results
                payload["raw_available_result_keys"] = sorted(raw_results.keys())
            return payload

        if force_refresh:
            _invalidate_project_cache_or_error(cache_service, identifier_id)
            _set_cache_header(response, "bypass")
            return _load_results()

        return _cache_get_or_set_payload(
            cache_service=cache_service,
            cache_key=cache_key,
            ttl_seconds=settings.XTJS_CACHE_PROJECT_RESULTS_TTL_SECONDS,
            response=response,
            factory=_load_results,
        )
    except HTTPException:
        raise
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"database error: {exc}") from exc


@router.get("/projects/{identifier_id}/results/{result_key}", summary="查询项目单项分析结果")
async def get_project_result_item(
    identifier_id: str,
    result_key: str,
    view: Literal["display", "raw"] = Query(default="display"),
    include_result_record: bool = Query(default=False),
    db_service: PostgreSQLService = Depends(get_db_service),
):
    """获取项目下特定键的分析结果。"""
    try:
        project = db_service.get_project_by_identifier(identifier_id)
        if not project:
            raise HTTPException(status_code=404, detail="project not found")

        result_record = db_service.get_project_result(identifier_id)
        refreshed_record, raw_results, display_results = _build_project_display_results(
            identifier_id=identifier_id,
            result_record=result_record,
            db_service=db_service,
        )
        selected_results = display_results if view == "display" else raw_results
        if view == "display":
            selected_results = _filter_visible_result_keys(selected_results)
        if result_key not in selected_results:
            raise HTTPException(status_code=404, detail="result_key not found")

        payload = {
            "result_record_meta": _build_result_record_meta(refreshed_record or result_record),
            "project": project,
            "result_key": result_key,
            "view": view,
            "result": selected_results[result_key],
            "available_result_keys": sorted(selected_results.keys()),
        }
        if include_result_record:
            payload["result_record"] = refreshed_record or result_record
        return payload
    except HTTPException:
        raise
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"database error: {exc}") from exc


@router.get("/projects/{identifier_id}/merged-results", summary="查询项目查重合并结果")
async def get_project_merged_results(
    identifier_id: str,
    result_key: Optional[str] = Query(default=None),
    include_result_record: bool = Query(default=False),
    db_service: PostgreSQLService = Depends(get_db_service),
):
    """获取查重后的合并结果（聚类视图）。"""
    try:
        project_detail = db_service.get_project_detail(identifier_id)
        if not project_detail:
            raise HTTPException(status_code=404, detail="project not found")
        project = project_detail.get("project") or {"identifier_id": identifier_id}

        requested_keys: Optional[list[str]] = None
        if result_key is not None:
            normalized_key = str(result_key or "").strip()
            if normalized_key not in DOC_TYPE_BY_MERGED_RESULT_KEY:
                raise HTTPException(status_code=400, detail="unsupported merged result key")
            requested_keys = [normalized_key]

        result_record = db_service.get_project_result(identifier_id)
        refreshed_record, merged_results = _load_or_build_project_merged_results(
            identifier_id=identifier_id,
            result_record=result_record,
            db_service=db_service,
            requested_keys=requested_keys,
        )
        if requested_keys and requested_keys[0] not in merged_results:
            raise HTTPException(status_code=404, detail="merged result not found")

        merged_result_keys = sorted(merged_results.keys())
        payload = {
            "project": project,
            "result_record_meta": _build_result_record_meta(refreshed_record or result_record),
            "results": merged_results,
            "available_result_keys": merged_result_keys,
        }
        if include_result_record:
            payload["result_record"] = refreshed_record or result_record
        return payload
    except HTTPException:
        raise
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"database error: {exc}") from exc


@router.get("/projects/{identifier_id}/visualization-data", summary="查询项目可视化聚合数据")
async def get_project_visualization_data(
    identifier_id: str,
    include_document_content: bool = Query(default=False),
    include_raw_results: bool = Query(default=False),
    include_result_record: bool = Query(default=False),
    db_service: PostgreSQLService = Depends(get_db_service),
):
    """返回前端可视化所需的聚合数据（项目、文档、关系、结果）。"""
    try:
        project_detail = db_service.get_project_detail(identifier_id)
        if not project_detail:
            raise HTTPException(status_code=404, detail="project not found")

        project_result = db_service.get_project_result(identifier_id)
        return _build_project_visualization_payload(
            identifier_id=identifier_id,
            project_detail=project_detail,
            project_result=project_result,
            db_service=db_service,
            include_document_content=include_document_content,
            include_raw_results=include_raw_results,
            include_result_record=include_result_record,
        )
    except HTTPException:
        raise
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"database error: {exc}") from exc


# 全局结果管理（不限定项目）
@router.get("/results", summary="查询结果表列表")
async def list_results(
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=20, ge=1, le=200),
    limit: Optional[int] = Query(default=None, ge=1, le=200),
    offset: Optional[int] = Query(default=None, ge=0),
    keyword: Optional[str] = Query(default=None),
    db_service: PostgreSQLService = Depends(get_db_service),
):
    """分页查看所有项目的结果记录。"""
    try:
        resolved_limit, resolved_offset = _resolve_pagination(
            page=page,
            page_size=page_size,
            limit=limit,
            offset=offset,
        )
        return db_service.list_project_results(
            limit=resolved_limit,
            offset=resolved_offset,
            keyword=keyword,
        )
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"database error: {exc}") from exc


@router.post("/results", summary="创建或覆盖项目结果")
async def create_or_replace_result(
    payload: ProjectResultUpsertRequest,
    db_service: PostgreSQLService = Depends(get_db_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """创建或完全替换某个项目的结果数据。"""
    try:
        result = db_service.create_or_replace_project_result(
            project_identifier_id=payload.project_identifier_id,
            result=payload.result,
        )
        _invalidate_project_cache_or_error(cache_service, payload.project_identifier_id)
        return result
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"database error: {exc}") from exc


@router.get("/results/{project_identifier_id}", summary="查询单个项目结果")
async def get_result_record(
    project_identifier_id: str,
    db_service: PostgreSQLService = Depends(get_db_service),
):
    """按项目标识获取结果记录。"""
    try:
        result_record = db_service.get_project_result(project_identifier_id)
        if not result_record:
            raise HTTPException(status_code=404, detail="result record not found")
        return result_record
    except HTTPException:
        raise
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"database error: {exc}") from exc


@router.put("/results/{project_identifier_id}", summary="更新单个项目结果")
async def update_result_record(
    project_identifier_id: str,
    payload: ProjectResultUpdateRequest,
    db_service: PostgreSQLService = Depends(get_db_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """更新指定项目的结果数据（覆盖）。"""
    try:
        result = db_service.create_or_replace_project_result(
            project_identifier_id=project_identifier_id,
            result=payload.result,
        )
        _invalidate_project_cache_or_error(cache_service, project_identifier_id)
        return result
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"database error: {exc}") from exc


@router.post("/projects/{project_identifier_id}/export-report", summary="导出当前展示结果 Word 报告")
async def export_project_result_report(
    project_identifier_id: str,
    payload: ProjectReportExportRequest,
    db_service: PostgreSQLService = Depends(get_db_service),
    oss_service: MinioService = Depends(get_oss_service),
):
    """基于本次请求的展示结果生成 Word 报告，不持久化前端删减副本。"""
    try:
        return await _export_result_word_report(
            project_identifier_id=project_identifier_id,
            result_payload=payload.result,
            db_service=db_service,
            oss_service=oss_service,
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"database error: {exc}") from exc


@router.post("/documents/{identifier_id}/preview/pages/{page}", summary="POST document page preview")
async def post_document_page_preview(
    identifier_id: str,
    page: int,
    payload: DocumentPreviewRequest,
    db_service: PostgreSQLService = Depends(get_db_service),
    oss_service: MinioService = Depends(get_oss_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """Use request body for preview and highlight parameters."""
    return await _build_document_page_preview_response(
        identifier_id=identifier_id,
        page=page,
        highlight=payload.highlight,
        highlight_bbox=payload.highlight_bbox,
        highlight_rects=payload.highlight_rects,
        highlight_coordinate_space=payload.highlight_coordinate_space or "auto",
        db_service=db_service,
        oss_service=oss_service,
        cache_service=cache_service,
    )


@router.delete("/results/{project_identifier_id}", summary="删除单个项目结果")
async def delete_result_record(
    project_identifier_id: str,
    db_service: PostgreSQLService = Depends(get_db_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """删除某个项目的所有分析结果。"""
    try:
        deleted = db_service.delete_project_result(project_identifier_id)
        if not deleted:
            raise HTTPException(status_code=404, detail="result record not found")
        _invalidate_project_cache_or_error(cache_service, project_identifier_id)
        return {"status": "deleted"}
    except HTTPException:
        raise
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"database error: {exc}") from exc


@router.post("/results/batch-delete", summary="批量删除项目结果")
async def batch_delete_result_records(
    payload: IdentifierBatchDeleteRequest,
    db_service: PostgreSQLService = Depends(get_db_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """批量删除多个项目的结果记录。"""
    try:
        deleted_count = db_service.delete_project_results(payload.identifier_ids)
        _invalidate_project_cache_or_error(cache_service)
        return {
            "requested_count": len(payload.identifier_ids),
            "deleted_count": deleted_count,
        }
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"database error: {exc}") from exc


# 业务检查接口（查重、形式审查、人员复用、错别字）
@router.post("/projects/duplicate-check", summary="项目商务标/技术标查重")
async def project_duplicate_check(
    identifier_id: str = Query(...),
    document_scope: DuplicateCheckScope = Query(default=DuplicateCheckScope.ALL),
    max_evidence_sections: int = Query(default=5, ge=1, le=20),
    max_pairs_per_type: int = Query(default=0, ge=0, le=500),
    db_service: PostgreSQLService = Depends(get_db_service),
    duplicate_check_service: DuplicateCheckService = Depends(get_duplicate_check_service),
):
    """执行项目的商务标/技术标内容查重，结果持久化。"""
    try:
        project = await run_in_threadpool(_refresh_project_or_404, db_service, identifier_id)
        _ensure_project_ocr_idle(project, analysis_name="查重分析")
        _ensure_project_analysis_status(
            project,
            required_status=_required_parsing_status_for_duplicate_scope(document_scope),
            analysis_name="查重分析",
        )
        return await run_in_threadpool(
            _run_project_duplicate_check,
            identifier_id=identifier_id,
            document_types=_document_types_from_scope(document_scope),
            max_evidence_sections=max_evidence_sections,
            max_pairs_per_type=max_pairs_per_type,
            result_key="duplicate_check",
            db_service=db_service,
            duplicate_check_service=duplicate_check_service,
        )
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.post("/projects/business-bid-format-review", summary="项目商务标形式审查")
async def project_business_bid_format_review(
    identifier_id: str = Query(...),
    db_service: PostgreSQLService = Depends(get_db_service),
):
    """对项目中的商务标进行格式合规性检查。"""
    review_service = UnifiedBusinessReviewService(db_service=db_service)
    try:
        project = await run_in_threadpool(_refresh_project_or_404, db_service, identifier_id)
        _ensure_project_ocr_idle(project, analysis_name="商务标形式审查")
        _ensure_project_analysis_status(
            project,
            required_status=PostgreSQLService.PARSING_STATUS_BUSINESS_OCR_COMPLETED,
            analysis_name="商务标形式审查",
        )
        result = await run_in_threadpool(
            review_service.persist_project_business_review,
            project_identifier=identifier_id,
            result_key=UnifiedBusinessReviewService.BUSINESS_RESULT_KEY,
        )
        _invalidate_project_cache_by_identifier(identifier_id)
        return result
    except HTTPException:
        raise
    except ValueError as exc:
        if "project not found" in str(exc).lower():
            raise HTTPException(status_code=404, detail="项目不存在") from exc
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.get("/projects/{identifier_id}/business-bid-format-review/editable", summary="List editable business-bid review values")
async def get_business_bid_format_review_editable(
    identifier_id: str,
    db_service: PostgreSQLService = Depends(get_db_service),
):
    try:
        result_record = await run_in_threadpool(db_service.get_project_result, identifier_id)
        review = _business_review_from_record(result_record)
        manual_payload = await run_in_threadpool(
            _business_manual_payload_for_project,
            identifier_id=identifier_id,
            db_service=db_service,
        )
        items = _build_business_format_editable_items(review, manual_payload)
        return {
            "project_identifier_id": identifier_id,
            "result_key": BUSINESS_FORMAT_RESULT_KEY,
            "review_content_inputs": manual_payload,
            "items": items,
            "item_count": len(items),
        }
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"database error: {exc}") from exc


@router.put("/projects/{identifier_id}/business-bid-format-review/manual-inputs", summary="Save editable business-bid review values")
async def save_business_bid_format_review_manual_inputs(
    identifier_id: str,
    payload: BusinessBidManualReviewInputsRequest,
    db_service: PostgreSQLService = Depends(get_db_service),
):
    try:
        raw_items = [item.model_dump() for item in payload.items]
        result_record = await run_in_threadpool(
            _save_business_manual_inputs,
            identifier_id=identifier_id,
            db_service=db_service,
            raw_items=raw_items,
            invalidate_project_cache=_invalidate_project_cache_by_identifier,
        )
        review = _business_review_from_record(result_record)
        manual_payload = await run_in_threadpool(
            _business_manual_payload_for_project,
            identifier_id=identifier_id,
            db_service=db_service,
        )
        items = _build_business_format_editable_items(review, manual_payload)
        return {
            "project_identifier_id": identifier_id,
            "result_key": BUSINESS_FORMAT_RESULT_KEY,
            "review_content_inputs": manual_payload,
            "items": items,
            "item_count": len(items),
            "result_record": result_record,
        }
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"database error: {exc}") from exc


@router.post("/projects/{identifier_id}/business-bid-format-review/manual-rerun", summary="Rerun business-bid review with manual values")
async def rerun_business_bid_format_review_with_manual_inputs(
    identifier_id: str,
    payload: Optional[BusinessBidManualReviewInputsRequest] = Body(default=None),
    db_service: PostgreSQLService = Depends(get_db_service),
):
    try:
        project = await run_in_threadpool(_refresh_project_or_404, db_service, identifier_id)
        _ensure_project_ocr_idle(project, analysis_name="business bid manual rerun")

        result_record = None
        if payload is not None and payload.items:
            result_record = await run_in_threadpool(
                _save_business_manual_inputs,
                identifier_id=identifier_id,
                db_service=db_service,
                raw_items=[item.model_dump() for item in payload.items],
                invalidate_project_cache=_invalidate_project_cache_by_identifier,
            )

        manual_payload = await run_in_threadpool(
            _business_manual_payload_for_project,
            identifier_id=identifier_id,
            db_service=db_service,
        )
        review_service = UnifiedBusinessReviewService(db_service=db_service)
        rerun_result = await run_in_threadpool(
            review_service.persist_project_business_review,
            project_identifier=identifier_id,
            result_key=BUSINESS_FORMAT_RESULT_KEY,
        )
        review = rerun_result["review"]
        corrected_review = _apply_manual_business_review_inputs(
            review,
            manual_payload,
        )
        result_record = await run_in_threadpool(
            db_service.update_project_manual_review_result,
            project_identifier_id=identifier_id,
            result_key=BUSINESS_FORMAT_RESULT_KEY,
            result_value=corrected_review,
        )
        _invalidate_project_cache_by_identifier(identifier_id)
        return {
            "project_identifier_id": identifier_id,
            "result_key": BUSINESS_FORMAT_RESULT_KEY,
            "review": corrected_review,
            "review_content_inputs": manual_payload,
            MANUAL_REVIEW_RESULTS_KEY: manual_review_results_from_record(result_record),
            "items": _build_business_format_editable_items(corrected_review, manual_payload),
            "base_review": review,
            "result_record": result_record,
        }
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"database error: {exc}") from exc


@router.post("/projects/deviation-check", summary="项目偏离表检查")
async def project_deviation_check(
    identifier_id: str = Query(...),
    db_service: PostgreSQLService = Depends(get_db_service),
):
    """在商务标和技术标中查找商务/技术偏离表并执行偏离表检查。"""
    try:
        project = await run_in_threadpool(_refresh_project_or_404, db_service, identifier_id)
        _ensure_project_ocr_idle(project, analysis_name="偏离表检查")
        _ensure_project_analysis_status(
            project,
            required_status=PostgreSQLService.PARSING_STATUS_TECHNICAL_OCR_COMPLETED,
            analysis_name="偏离表检查",
        )
        return await run_in_threadpool(
            _run_project_deviation_check,
            identifier_id=identifier_id,
            db_service=db_service,
        )
    except HTTPException:
        raise
    except ValueError as exc:
        if "project not found" in str(exc).lower():
            raise HTTPException(status_code=404, detail="项目不存在") from exc
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.post("/projects/business-bid-duplicate-check", summary="项目商务标内容查重")
async def project_business_bid_duplicate_check(
    identifier_id: str = Query(...),
    max_evidence_sections: int = Query(default=5, ge=1, le=20),
    max_pairs_per_type: int = Query(default=0, ge=0, le=500),
    db_service: PostgreSQLService = Depends(get_db_service),
    duplicate_check_service: DuplicateCheckService = Depends(get_duplicate_check_service),
):
    """仅对商务标进行内容查重。"""
    try:
        project = await run_in_threadpool(_refresh_project_or_404, db_service, identifier_id)
        _ensure_project_ocr_idle(project, analysis_name="商务标查重")
        _ensure_project_analysis_status(
            project,
            required_status=PostgreSQLService.PARSING_STATUS_BUSINESS_OCR_COMPLETED,
            analysis_name="商务标查重",
        )
        return await run_in_threadpool(
            _run_project_duplicate_check,
            identifier_id=identifier_id,
            document_types=[DuplicateCheckScope.BUSINESS_BID.value],
            max_evidence_sections=max_evidence_sections,
            max_pairs_per_type=max_pairs_per_type,
            result_key="business_bid_duplicate_check",
            db_service=db_service,
            duplicate_check_service=duplicate_check_service,
        )
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.post("/projects/business-bid-duplicate-check/upload-json", summary="上传 OCR JSON 并执行商务标内容查重")
async def upload_business_bid_duplicate_check(
    tender_json_file: UploadFile = File(...),
    business_bid_json_files: list[UploadFile] = File(...),
    technical_bid_json_files: Optional[list[UploadFile]] = File(default=None),
    project_name: Optional[str] = Form(default=None, description="项目名称；不传时自动生成临时项目名"),
    max_evidence_sections: int = Form(default=5, ge=1, le=20),
    max_pairs_per_type: int = Form(default=0, ge=0, le=500),
    db_service: PostgreSQLService = Depends(get_db_service),
    duplicate_check_service: DuplicateCheckService = Depends(get_duplicate_check_service),
):
    """上传招投标 OCR JSON 文件，直接执行商务标内容查重。"""
    uploads = [upload for upload in business_bid_json_files if upload is not None]
    if not uploads:
        raise HTTPException(status_code=400, detail="business_bid_json_files 不能为空。")

    try:
        persisted_documents, document_records = await _persist_uploaded_analysis_documents(
            tender_json_file=tender_json_file,
            business_bid_json_files=uploads,
            technical_bid_json_files=technical_bid_json_files,
            project_name=project_name,
            db_service=db_service,
        )
        resolved_project_identifier = persisted_documents["project"]["identifier_id"]
        duplicate_result = await run_in_threadpool(
            duplicate_check_service.check_project_documents,
            project_identifier=resolved_project_identifier,
            project=_project_snapshot(resolved_project_identifier),
            document_records=document_records,
            document_types=[DuplicateCheckScope.BUSINESS_BID.value],
            max_evidence_sections=max_evidence_sections,
            max_pairs_per_type=max_pairs_per_type,
        )
        result_record = await run_in_threadpool(
            _persist_uploaded_result,
            db_service=db_service,
            project_identifier=resolved_project_identifier,
            result_key="business_bid_duplicate_check",
            result_value=duplicate_result,
        )
        merged_results = await run_in_threadpool(
            _persist_duplicate_merge_results,
            db_service=db_service,
            project_identifier=resolved_project_identifier,
            source_result_key="business_bid_duplicate_check",
            raw_result=duplicate_result,
        )
        return {
            "project": persisted_documents["project"],
            "result_key": "business_bid_duplicate_check",
            "result": duplicate_result,
            "result_record": result_record,
            "merged_results": merged_results,
            "document_binding": persisted_documents["binding"],
        }
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.post("/projects/technical-bid-duplicate-check", summary="项目技术标内容查重")
async def project_technical_bid_duplicate_check(
    identifier_id: str = Query(...),
    max_evidence_sections: int = Query(default=5, ge=1, le=20),
    max_pairs_per_type: int = Query(default=0, ge=0, le=500),
    db_service: PostgreSQLService = Depends(get_db_service),
    duplicate_check_service: DuplicateCheckService = Depends(get_duplicate_check_service),
):
    """仅对技术标进行内容查重。"""
    try:
        project = await run_in_threadpool(_refresh_project_or_404, db_service, identifier_id)
        _ensure_project_ocr_idle(project, analysis_name="技术标查重")
        _ensure_project_analysis_status(
            project,
            required_status=PostgreSQLService.PARSING_STATUS_TECHNICAL_OCR_COMPLETED,
            analysis_name="技术标查重",
        )
        return await run_in_threadpool(
            _run_project_duplicate_check,
            identifier_id=identifier_id,
            document_types=[DuplicateCheckScope.TECHNICAL_BID.value],
            max_evidence_sections=max_evidence_sections,
            max_pairs_per_type=max_pairs_per_type,
            result_key="technical_bid_duplicate_check",
            db_service=db_service,
            duplicate_check_service=duplicate_check_service,
        )
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.post("/projects/technical-bid-duplicate-check/upload-json", summary="上传 OCR JSON 并执行技术标内容查重")
async def upload_technical_bid_duplicate_check(
    tender_json_file: UploadFile = File(...),
    technical_bid_json_files: list[UploadFile] = File(...),
    business_bid_json_files: Optional[list[UploadFile]] = File(default=None),
    project_name: Optional[str] = Form(default=None, description="项目名称；不传时自动生成临时项目名"),
    max_evidence_sections: int = Form(default=5, ge=1, le=20),
    max_pairs_per_type: int = Form(default=0, ge=0, le=500),
    db_service: PostgreSQLService = Depends(get_db_service),
    duplicate_check_service: DuplicateCheckService = Depends(get_duplicate_check_service),
):
    """上传招投标 OCR JSON 文件，直接执行技术标内容查重。"""
    uploads = [upload for upload in technical_bid_json_files if upload is not None]
    if not uploads:
        raise HTTPException(status_code=400, detail="technical_bid_json_files 不能为空。")

    try:
        persisted_documents, document_records = await _persist_uploaded_analysis_documents(
            tender_json_file=tender_json_file,
            business_bid_json_files=business_bid_json_files,
            technical_bid_json_files=uploads,
            project_name=project_name,
            db_service=db_service,
        )
        resolved_project_identifier = persisted_documents["project"]["identifier_id"]
        duplicate_result = await run_in_threadpool(
            duplicate_check_service.check_project_documents,
            project_identifier=resolved_project_identifier,
            project=_project_snapshot(resolved_project_identifier),
            document_records=document_records,
            document_types=[DuplicateCheckScope.TECHNICAL_BID.value],
            max_evidence_sections=max_evidence_sections,
            max_pairs_per_type=max_pairs_per_type,
        )
        result_record = await run_in_threadpool(
            _persist_uploaded_result,
            db_service=db_service,
            project_identifier=resolved_project_identifier,
            result_key="technical_bid_duplicate_check",
            result_value=duplicate_result,
        )
        merged_results = await run_in_threadpool(
            _persist_duplicate_merge_results,
            db_service=db_service,
            project_identifier=resolved_project_identifier,
            source_result_key="technical_bid_duplicate_check",
            raw_result=duplicate_result,
        )
        return {
            "project": persisted_documents["project"],
            "result_key": "technical_bid_duplicate_check",
            "result": duplicate_result,
            "result_record": result_record,
            "merged_results": merged_results,
            "document_binding": persisted_documents["binding"],
        }
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.post("/projects/personnel-reuse-check", summary="项目一人多用检查")
async def project_personnel_reuse_check(
    identifier_id: str = Query(...),
    payload: PersonnelReuseCheckRequest | None = Body(default=None),
    db_service: PostgreSQLService = Depends(get_db_service),
    bid_document_review_service: BidDocumentReviewService = Depends(get_bid_document_review_service),
):
    """抽取商务标/技术标人员，支持传入确认后名单再做一人多用检查。"""
    try:
        project = await run_in_threadpool(_refresh_project_or_404, db_service, identifier_id)
        _ensure_project_ocr_idle(project, analysis_name="一人多用检查")
        _ensure_project_analysis_status(
            project,
            required_status=PostgreSQLService.PARSING_STATUS_TECHNICAL_OCR_COMPLETED,
            analysis_name="一人多用检查",
        )
        return await run_in_threadpool(
            _run_project_personnel_reuse_check,
            identifier_id=identifier_id,
            db_service=db_service,
            bid_document_review_service=bid_document_review_service,
            confirmed_names=(payload.confirmed_names if payload else None),
        )
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.put("/projects/{identifier_id}/personnel-reuse-draft", summary="保存项目人员复用草稿")
async def save_project_personnel_reuse_draft(
    identifier_id: str,
    payload: PersonnelReuseDraftRequest,
    db_service: PostgreSQLService = Depends(get_db_service),
    bid_document_review_service: BidDocumentReviewService = Depends(get_bid_document_review_service),
):
    """保存结果审核页按文件编辑的人员草稿，不生成跨文件重名问题。"""
    try:
        project = await run_in_threadpool(_refresh_project_or_404, db_service, identifier_id)
        _ensure_project_ocr_idle(project, analysis_name="保存人员复用草稿")
        return await run_in_threadpool(
            _persist_project_personnel_reuse_draft,
            identifier_id=identifier_id,
            db_service=db_service,
            bid_document_review_service=bid_document_review_service,
            draft_documents=[item.model_dump() for item in payload.documents],
            confirmation_status="draft",
        )
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.post("/projects/{identifier_id}/personnel-reuse-confirm", summary="确认项目人员并重算复用")
async def confirm_project_personnel_reuse_draft(
    identifier_id: str,
    payload: PersonnelReuseDraftRequest,
    db_service: PostgreSQLService = Depends(get_db_service),
    bid_document_review_service: BidDocumentReviewService = Depends(get_bid_document_review_service),
):
    """保存人员草稿，并按草稿重算全部投标文件之间的重名复用结果。"""
    try:
        project = await run_in_threadpool(_refresh_project_or_404, db_service, identifier_id)
        _ensure_project_ocr_idle(project, analysis_name="确认人员复用")
        return await run_in_threadpool(
            _persist_project_personnel_reuse_draft,
            identifier_id=identifier_id,
            db_service=db_service,
            bid_document_review_service=bid_document_review_service,
            draft_documents=[item.model_dump() for item in payload.documents],
            confirmation_status="confirmed",
        )
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.post("/projects/personnel-reuse-check/upload-json", summary="上传 OCR JSON 并执行一人多用检查")
async def upload_personnel_reuse_check(
    tender_json_file: UploadFile = File(...),
    business_bid_json_files: Optional[list[UploadFile]] = File(default=None),
    technical_bid_json_files: Optional[list[UploadFile]] = File(default=None),
    project_name: Optional[str] = Form(default=None, description="项目名称；不传时自动生成临时项目名"),
    confirmed_names_json: Optional[str] = Form(default=None, description="业务确认后的人名 JSON 数组。"),
    db_service: PostgreSQLService = Depends(get_db_service),
    bid_document_review_service: BidDocumentReviewService = Depends(get_bid_document_review_service),
):
    """上传招投标 OCR JSON 文件，直接执行人员复用检查。"""
    business_uploads = [upload for upload in (business_bid_json_files or []) if upload is not None]
    technical_uploads = [upload for upload in (technical_bid_json_files or []) if upload is not None]
    if not business_uploads and not technical_uploads:
        raise HTTPException(status_code=400, detail="至少需要上传一份商务标或技术标文件。")
    confirmed_names = None
    if confirmed_names_json and confirmed_names_json.strip():
        try:
            confirmed_names = json.loads(confirmed_names_json)
        except json.JSONDecodeError as exc:
            raise HTTPException(status_code=400, detail="confirmed_names_json 必须是 JSON 数组。") from exc
        if not isinstance(confirmed_names, list):
            raise HTTPException(status_code=400, detail="confirmed_names_json 必须是 JSON 数组。")

    try:
        persisted_documents, document_records = await _persist_uploaded_analysis_documents(
            tender_json_file=tender_json_file,
            business_bid_json_files=business_uploads,
            technical_bid_json_files=technical_uploads,
            project_name=project_name,
            db_service=db_service,
        )
        resolved_project_identifier = persisted_documents["project"]["identifier_id"]
        personnel_result = await run_in_threadpool(
            bid_document_review_service.check_project_personnel_reuse,
            project_identifier=resolved_project_identifier,
            project=_project_snapshot(resolved_project_identifier),
            document_records=document_records,
            document_types=None,
            confirmed_names=confirmed_names,
        )
        result_record = await run_in_threadpool(
            _persist_uploaded_result,
            db_service=db_service,
            project_identifier=resolved_project_identifier,
            result_key="personnel_reuse_check",
            result_value=personnel_result,
        )
        return {
            "project": persisted_documents["project"],
            "result_key": "personnel_reuse_check",
            "result": personnel_result,
            "result_record": result_record,
            "document_binding": persisted_documents["binding"],
        }
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc




@router.post("/projects/{identifier_id}/duplicate-check", summary="项目商务标/技术标查重", include_in_schema=False)
async def project_duplicate_check_legacy(
    identifier_id: str,
    payload: Optional[ProjectDuplicateCheckRequest] = None,
    db_service: PostgreSQLService = Depends(get_db_service),
    duplicate_check_service: DuplicateCheckService = Depends(get_duplicate_check_service),
):
    request_payload = payload or ProjectDuplicateCheckRequest()
    project = await run_in_threadpool(_refresh_project_or_404, db_service, identifier_id)
    _ensure_project_ocr_idle(project, analysis_name="查重分析")
    _ensure_project_analysis_status(
        project,
        required_status=_required_parsing_status_for_document_types(request_payload.document_types),
        analysis_name="查重分析",
    )
    return await run_in_threadpool(
        _run_project_duplicate_check,
        identifier_id=identifier_id,
        document_types=request_payload.document_types,
        max_evidence_sections=request_payload.max_evidence_sections,
        max_pairs_per_type=request_payload.max_pairs_per_type,
        result_key="duplicate_check",
        db_service=db_service,
        duplicate_check_service=duplicate_check_service,
    )


# 文档与绑定关系管理
@router.post("/projects/{identifier_id}/bind-documents", summary="绑定招标/商务标/技术标文件")
async def bind_project_documents(
    identifier_id: str,
    payload: ProjectBindDocumentsRequest,
    db_service: PostgreSQLService = Depends(get_db_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """将三个文档标识绑定到项目下形成一个关系记录。"""
    try:
        relation = db_service.bind_project_documents(
            identifier_id,
            payload.tender_document_identifier,
            payload.business_bid_document_identifier,
            payload.technical_bid_document_identifier,
        )
        _invalidate_project_cache_or_error(cache_service, str(relation.get("project_identifier") or identifier_id))
        return relation
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.get("/relations", summary="查询项目文件绑定列表")
async def list_relations(
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=20, ge=1, le=200),
    limit: Optional[int] = Query(default=None, ge=1, le=200),
    offset: Optional[int] = Query(default=None, ge=0),
    keyword: Optional[str] = Query(default=None),
    project_identifier: Optional[str] = Query(default=None),
    db_service: PostgreSQLService = Depends(get_db_service),
):
    """分页查询项目文档绑定关系。"""
    try:
        resolved_limit, resolved_offset = _resolve_pagination(
            page=page,
            page_size=page_size,
            limit=limit,
            offset=offset,
        )
        return db_service.list_relations(
            limit=resolved_limit,
            offset=resolved_offset,
            keyword=keyword,
            project_identifier=project_identifier,
        )
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.get("/relations/{relation_id}", summary="查询关联详情")
async def get_relation_detail(
    relation_id: int,
    db_service: PostgreSQLService = Depends(get_db_service),
):
    """按 ID 获取单条绑定关系。"""
    try:
        relation = db_service.get_relation_by_id(relation_id)
        if not relation:
            raise HTTPException(status_code=404, detail="关联不存在")
        return relation
    except HTTPException:
        raise
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.put("/relations/{relation_id}", summary="更新关联")
async def update_relation(
    relation_id: int,
    payload: ProjectRelationUpdateRequest,
    db_service: PostgreSQLService = Depends(get_db_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """修改已有的文档绑定关系。"""
    try:
        updated = db_service.update_relation(
            relation_id=relation_id,
            tender_document_identifier=payload.tender_document_identifier,
            business_bid_document_identifier=payload.business_bid_document_identifier,
            technical_bid_document_identifier=payload.technical_bid_document_identifier,
        )
        if not updated:
            raise HTTPException(status_code=404, detail="关联不存在")
        _invalidate_project_cache_or_error(cache_service, str(updated.get("project_identifier") or ""))
        return updated
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.delete("/relations/{relation_id}", summary="删除关联")
async def delete_relation(
    relation_id: int,
    db_service: PostgreSQLService = Depends(get_db_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """删除一条文档绑定关系。"""
    try:
        deleted = db_service.delete_relation(relation_id)
        if not deleted:
            raise HTTPException(status_code=404, detail="关联不存在")
        _invalidate_project_cache_or_error(cache_service)
        return {"status": "deleted"}
    except HTTPException:
        raise
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.post("/relations/batch-delete", summary="批量删除关联")
async def batch_delete_relations(
    payload: RelationBatchDeleteRequest,
    db_service: PostgreSQLService = Depends(get_db_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """批量删除绑定关系。"""
    try:
        deleted_count = db_service.delete_relations(payload.relation_ids)
        _invalidate_project_cache_or_error(cache_service)
        return {
            "requested_count": len(payload.relation_ids),
            "deleted_count": deleted_count,
        }
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


# 文档 CUD 与预览
@router.post("/documents", summary="上传并创建文档")
async def create_document(
    file: UploadFile = File(...),
    document_type: DocumentType = Form(...),
    identifier_id: Optional[str] = Form(default=None),
    document_name: Optional[str] = Form(default=None),
    object_name: Optional[str] = Form(default=None),
    recognition_options: RecognitionOptions = Depends(get_form_recognition_options),
    db_service: PostgreSQLService = Depends(get_db_service),
    oss_service: MinioService = Depends(get_oss_service),
    analysis_service=Depends(get_text_analysis_service),
):
    """上传文件并触发 OCR 提取，创建文档记录。"""
    result = await upload_extract_and_create_document(
        file=file,
        document_type=document_type,
        db_service=db_service,
        oss_service=oss_service,
        analysis_service=analysis_service,
        identifier_id=identifier_id,
        document_name=document_name,
        object_name=object_name,
        **recognition_options.as_kwargs(),
        raise_http_exception=True,
    )
    return {
        "document": result["document"],
        "upload": result["upload"],
    }


@router.get("/documents", summary="查询文档列表")
async def list_documents(
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=20, ge=1, le=200),
    limit: Optional[int] = Query(default=None, ge=1, le=200),
    offset: Optional[int] = Query(default=None, ge=0),
    keyword: Optional[str] = Query(default=None),
    document_type: Optional[str] = Query(default=None),
    extracted: Optional[bool] = Query(default=None),
    db_service: PostgreSQLService = Depends(get_db_service),
):
    """分页查询文档列表，支持按类型、提取状态过滤。"""
    try:
        resolved_limit, resolved_offset = _resolve_pagination(
            page=page,
            page_size=page_size,
            limit=limit,
            offset=offset,
        )
        return db_service.list_documents(
            limit=resolved_limit,
            offset=resolved_offset,
            keyword=keyword,
            document_type=document_type,
            extracted=extracted,
        )
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.post("/documents/batch-delete", summary="批量删除文档")
async def batch_delete_documents(
    payload: IdentifierBatchDeleteRequest,
    db_service: PostgreSQLService = Depends(get_db_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """软删除一批文档。"""
    try:
        deleted_count = db_service.soft_delete_documents(payload.identifier_ids)
        _invalidate_project_cache_or_error(cache_service)
        for document_identifier in payload.identifier_ids:
            if str(document_identifier or "").strip():
                _invalidate_document_preview_cache_or_error(cache_service, document_identifier)
        return {
            "requested_count": len(payload.identifier_ids),
            "deleted_count": deleted_count,
        }
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.get("/documents/{identifier_id}", summary="查询文档")
async def get_document(
    identifier_id: str,
    db_service: PostgreSQLService = Depends(get_db_service),
):
    """按标识获取单个文档详情。"""
    try:
        document = db_service.get_document_by_identifier(identifier_id)
        if not document:
            raise HTTPException(status_code=404, detail="文档不存在")
        return document
    except HTTPException:
        raise
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.get("/documents/{identifier_id}/review-content", summary="查询文档人工识别工作副本")
async def get_document_review_content(
    identifier_id: str,
    db_service: PostgreSQLService = Depends(get_db_service),
):
    try:
        payload = db_service.get_document_review_content(identifier_id)
        if not payload:
            raise HTTPException(status_code=404, detail="文档不存在")
        return payload
    except HTTPException:
        raise
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.put("/documents/{identifier_id}/review-content", summary="保存文档人工识别工作副本")
async def update_document_review_content(
    identifier_id: str,
    payload: DocumentReviewContentUpdateRequest,
    db_service: PostgreSQLService = Depends(get_db_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    try:
        updated = db_service.update_document_review_content(
            identifier_id,
            effective_content=payload.effective_content,
            inputs=payload.inputs,
        )
        if not updated:
            raise HTTPException(status_code=404, detail="文档不存在")
        _invalidate_project_cache_or_error(cache_service)
        _invalidate_document_preview_cache_or_error(cache_service, identifier_id)
        return updated
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.put("/documents/{identifier_id}", summary="更新文档")
async def update_document(
    identifier_id: str,
    payload: DocumentUpdateRequest,
    db_service: PostgreSQLService = Depends(get_db_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """更新文档的文件名或文件 URL。"""
    try:
        normalized_file_url = (
            normalize_file_url(payload.file_url) if payload.file_url is not None else None
        )
        updated = db_service.update_document(
            identifier_id=identifier_id,
            file_name=payload.file_name,
            file_url=normalized_file_url,
        )
        if not updated:
            raise HTTPException(status_code=404, detail="文档不存在")
        _invalidate_project_cache_or_error(cache_service)
        _invalidate_document_preview_cache_or_error(cache_service, str(updated.get("identifier_id") or identifier_id))
        return updated
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.delete("/documents/{identifier_id}", summary="删除文档")
async def delete_document(
    identifier_id: str,
    db_service: PostgreSQLService = Depends(get_db_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """软删除文档。"""
    try:
        deleted = db_service.soft_delete_document(identifier_id)
        if not deleted:
            raise HTTPException(status_code=404, detail="文档不存在")
        _invalidate_project_cache_or_error(cache_service)
        _invalidate_document_preview_cache_or_error(cache_service, identifier_id)
        return {"status": "deleted"}
    except HTTPException:
        raise
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"数据库错误：{exc}") from exc


@router.get("/documents/{identifier_id}/source", summary="获取文档源文件")
async def get_document_source(
    identifier_id: str,
    page: Optional[int] = Query(default=None, ge=1, description="可选页码，仅 PDF 源文件支持跳转"),
    db_service: PostgreSQLService = Depends(get_db_service),
    oss_service: MinioService = Depends(get_oss_service),
):
    """重定向到文档在 MinIO 中的预签名下载 URL，支持 PDF 页码锚点。"""
    try:
        document = db_service.get_document_by_identifier(identifier_id)
        if not document:
            raise HTTPException(status_code=404, detail="document not found")

        bucket_name, object_name, _file_name = _resolve_document_source_object(document)
        presigned_url = oss_service.get_presigned_url(object_name, bucket_name)
        if page and _document_source_kind(document) == "pdf":
            presigned_url = f"{presigned_url}#page={page}"
        return RedirectResponse(url=presigned_url, status_code=307)
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"database error: {exc}") from exc


async def _build_document_page_preview_response(
    identifier_id: str,
    page: int,
    highlight: Optional[list[str]] = Query(default=None, description="需要高亮的关键词，可传多个"),
    highlight_bbox: Optional[str] = Query(default=None, description="单个高亮框，格式为逗号分隔四个数字"),
    highlight_rects: Optional[str] = Query(default=None, description="多个高亮框的 JSON 数组字符串"),
    highlight_coordinate_space: str = Query(default="auto", description="高亮坐标系：auto/pdf_point/pdf/ocr_image"),
    db_service: PostgreSQLService = Depends(get_db_service),
    oss_service: MinioService = Depends(get_oss_service),
    cache_service: RedisCacheService = Depends(get_cache_service),
):
    """返回文档指定页的 base64 预览图，支持文本/区域高亮。结果会被缓存。"""
    try:
        document = db_service.get_document_by_identifier(identifier_id)
        if not document:
            raise HTTPException(status_code=404, detail="document not found")

        normalized_highlight_phrases = _normalize_preview_highlight_phrases(highlight)
        normalized_highlight_bbox = _normalize_preview_highlight_bbox(highlight_bbox)
        normalized_highlight_rects = _normalize_preview_highlight_rects(highlight_rects)
        resolved_coordinate_space = _document_preview_coordinate_space(
            document,
            highlight_coordinate_space,
        )
        has_highlight_payload = bool(
            normalized_highlight_phrases or
            normalized_highlight_bbox or
            normalized_highlight_rects
        )
        source_kind = _document_source_kind(document)
        if not has_highlight_payload:
            payload = _load_or_create_raw_preview_payload(
                document=document,
                page=page,
                source_kind=source_kind,
                cache_service=cache_service,
                oss_service=oss_service,
            )
            payload["document_identifier"] = identifier_id
            payload["file_name"] = str(document.get("file_name") or "")
            payload["source_url"] = f"/api/postgresql/documents/{identifier_id}/source"
            return JSONResponse(payload)

        file_bytes, _content_type, _object_name = _load_document_source_bytes(
            document=document,
            oss_service=oss_service,
        )
        try:
            payload = _preview_payload_from_source(
                file_bytes=file_bytes,
                source_kind=source_kind,
                page=page,
                highlight_phrases=normalized_highlight_phrases,
                highlight_bbox=normalized_highlight_bbox,
                highlight_rects=normalized_highlight_rects,
                document=document,
                highlight_coordinate_space=resolved_coordinate_space,
            )
        except Exception:
            # 高亮渲染失败时回退到无高亮预览
            payload = _preview_payload_from_source(
                file_bytes=file_bytes,
                source_kind=source_kind,
                page=page,
                highlight_phrases=[],
                highlight_bbox=None,
                highlight_rects=[],
                document=document,
                highlight_coordinate_space=resolved_coordinate_space,
            )
            payload["highlight_applied"] = False
            payload["highlight_fallback"] = True
        payload["document_identifier"] = identifier_id
        payload["file_name"] = str(document.get("file_name") or "")
        payload["source_url"] = f"/api/postgresql/documents/{identifier_id}/source"
        return JSONResponse(payload)
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except CacheUnavailableError as exc:
        raise HTTPException(status_code=503, detail=f"预览加载失败，请检查 Redis：{exc}") from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=f"预览加载失败，请检查 MinIO：{exc}") from exc
    except PsycopgError as exc:
        raise HTTPException(status_code=500, detail=f"database error: {exc}") from exc    # ???????
